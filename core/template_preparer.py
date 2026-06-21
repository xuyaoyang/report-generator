"""
Analyze template structure and build a replacement mapping.
Strategy: Group consecutive yellow runs as logical fields,
then assign meaningful placeholder names based on position context.
"""
import json, os, re
from docx import Document
from docx.oxml.ns import qn


def is_yellow(run):
    """Check if a run has yellow highlighting."""
    if run.font.highlight_color is not None:
        hl = str(run.font.highlight_color)
        return 'YELLOW' in hl.upper() or hl == '7'
    return False


def clear_highlight(run):
    rPr = run._element.get_or_add_rPr()
    h = rPr.find(qn('w:highlight'))
    if h is not None:
        rPr.remove(h)


def extract_yellow_fields(doc):
    """
    Extract all yellow-highlighted fields, grouping consecutive yellow runs.
    Returns list of field descriptors.
    """
    fields = []
    field_id = 0

    def extract_from_paragraphs(paragraphs, location_prefix, extra_info=None):
        nonlocal field_id
        for p_idx, para in enumerate(paragraphs):
            text = para.text.strip()
            if not text:
                continue
            # Group consecutive yellow runs
            yellow_groups = []
            current_group = []
            for run in para.runs:
                if is_yellow(run):
                    current_group.append(run)
                else:
                    if current_group:
                        yellow_groups.append(current_group)
                        current_group = []
            if current_group:
                yellow_groups.append(current_group)

            for group in yellow_groups:
                combined_text = ''.join(r.text for r in group)
                if combined_text.strip():
                    field_id += 1
                    location = f'{location_prefix}[{p_idx}]'
                    if extra_info:
                        location += f' ({extra_info})'
                    fields.append({
                        'id': field_id,
                        'placeholder': f'{{{{FIELD_{field_id:03d}}}}}',
                        'original': combined_text.strip(),
                        'location': location,
                        'runs': group,  # reference to actual runs for replacement
                        'para': para,
                    })
        return fields

    # Body paragraphs
    for p_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        yellow_groups = []
        current_group = []
        for run in para.runs:
            if is_yellow(run):
                current_group.append(run)
            else:
                if current_group:
                    yellow_groups.append(current_group)
                    current_group = []
        if current_group:
            yellow_groups.append(current_group)

        for group in yellow_groups:
            combined_text = ''.join(r.text for r in group)
            if combined_text.strip():
                field_id += 1
                fields.append({
                    'id': field_id,
                    'placeholder': f'{{{{FIELD_{field_id:03d}}}}}',
                    'original': combined_text.strip(),
                    'location': f'paragraph[{p_idx}]',
                    'para': para,
                    'runs': group,
                })

    # Table cells — deduplicate merged cells that share the same XML runs
    seen_run_ids = set()

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    yellow_groups = []
                    current_group = []
                    for run in para.runs:
                        if is_yellow(run):
                            current_group.append(run)
                        else:
                            if current_group:
                                yellow_groups.append(current_group)
                                current_group = []
                    if current_group:
                        yellow_groups.append(current_group)

                    for group in yellow_groups:
                        combined_text = ''.join(r.text for r in group)
                        if not combined_text.strip():
                            continue
                        # Deduplicate: merged cells share the same XML runs.
                        # Use the first run's element id as a proxy for uniqueness.
                        run_key = id(group[0]._element)
                        if run_key in seen_run_ids:
                            continue
                        seen_run_ids.add(run_key)

                        field_id += 1
                        fields.append({
                            'id': field_id,
                            'placeholder': f'{{{{FIELD_{field_id:03d}}}}}',
                            'original': combined_text.strip(),
                            'location': f'table[{t_idx}].cell[{r_idx},{c_idx}]',
                            'para': para,
                            'runs': group,
                        })

    return fields


def assign_semantic_names(fields):
    """
    Assign meaningful names to fields based on position and original content.
    Returns categorized fields.
    """
    categories = {
        'cover_info': [],      # 封面信息
        'certificates': [],    # 合格证数据
        'mechanical': [],      # 力学性能数据
        'visual': [],          # 外观尺寸数据
        'unknown': [],         # 未分类
    }

    for f in fields:
        loc = f['location']
        orig = f['original']

        # Cover page paragraphs (first few paragraphs in document)
        if loc.startswith('paragraph['):
            p_idx = int(re.search(r'\[(\d+)\]', loc).group(1))
            if p_idx <= 30:
                categories['cover_info'].append(f)
                continue

        # Certificate tables (tables 0-5 based on analysis)
        table_match = re.search(r'table\[(\d+)\]', loc)
        if table_match:
            t_idx = int(table_match.group(1))
            if t_idx <= 5:
                categories['certificates'].append(f)
            elif t_idx <= 8:
                categories['mechanical'].append(f)
            else:
                categories['visual'].append(f)
        else:
            categories['unknown'].append(f)

    return categories


def prepare_template(input_path, output_path, mapping_path, add_page_breaks=True):
    doc = Document(input_path)

    # Extract fields
    fields = extract_yellow_fields(doc)
    print(f'Total yellow fields found: {len(fields)}')

    # Categorize
    cats = assign_semantic_names(fields)
    for cat, items in cats.items():
        print(f'  {cat}: {len(items)} fields')

    # Replace all yellow runs with placeholders
    for f in fields:
        # Replace first run's text with placeholder, clear the rest
        f['runs'][0].text = f['placeholder']
        clear_highlight(f['runs'][0])
        for run in f['runs'][1:]:
            run.text = ''
            clear_highlight(run)

    # Save
    doc.save(output_path)

    # Build mapping JSON
    mapping = {
        'total_fields': len(fields),
        'categories': {},
    }
    for cat, items in cats.items():
        mapping['categories'][cat] = []
        for f in items:
            mapping['categories'][cat].append({
                'placeholder': f['placeholder'],
                'original': f['original'],
                'location': f['location'],
            })

    with open(mapping_path, 'w', encoding='utf-8') as fp:
        json.dump(mapping, fp, ensure_ascii=False, indent=2)

    # Add page breaks at key positions
    if add_page_breaks:
        from core.template_layout_fixer import fix_template_pagination
        fix_template_pagination(output_path)

    print(f'\nTemplate saved: {output_path}')
    print(f'Mapping saved: {mapping_path}')

    # Print cover info fields
    print('\n=== Cover Info Fields ===')
    for f in cats['cover_info']:
        print(f"  {f['placeholder']} = '{f['original']}' @ {f['location']}")

    print('\n=== Certificate Fields (first 10) ===')
    for f in cats['certificates'][:10]:
        print(f"  {f['placeholder']} = '{f['original'][:60]}' @ {f['location']}")

    return mapping


if __name__ == '__main__':
    _base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    input_docx = os.path.join(_base, 'products', 'isolation_bearing', 'template.docx')
    output_docx = os.path.join(_base, 'products', 'isolation_bearing', 'template_prepared.docx')
    mapping_json = os.path.join(_base, 'products', 'isolation_bearing', 'param_mapping.json')

    prepare_template(input_docx, output_docx, mapping_json)
