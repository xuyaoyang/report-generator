"""
Convert .doc template to .docx and analyze structure.
Uses Word COM to convert, then python-docx to analyze.
"""
import os
import json

# Project root (2 levels up from core/)
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE_SRC = os.path.join(BASE_DIR, '隔震支座出厂报告2025-2-7.doc')
TEMPLATE_DOCX = os.path.join(BASE_DIR, 'products', 'isolation_bearing', 'template.docx')


def convert_doc_to_docx(doc_path, docx_path):
    """Convert .doc to .docx using Word COM."""
    import win32com.client
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    try:
        doc = word.Documents.Open(doc_path)
        doc.SaveAs2(docx_path, FileFormat=16)  # 16 = wdFormatDocumentDefault
        doc.Close()
        return True
    finally:
        word.Quit()


def analyze_template(docx_path):
    """Analyze .docx structure: paragraphs with highlights, tables."""
    from docx import Document

    doc = Document(docx_path)

    result = {
        'paragraphs': [],
        'tables': [],
        'sections': len(doc.sections),
    }

    # Analyze paragraphs
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        info = {
            'index': i,
            'text': text[:200],
            'style': para.style.name if para.style else '',
        }
        # Check for highlighted runs
        highlights = set()
        for run in para.runs:
            if run.font.highlight_color is not None:
                highlights.add(str(run.font.highlight_color))
        info['highlights'] = list(highlights) if highlights else []
        result['paragraphs'].append(info)

    # Analyze tables
    for t_idx, table in enumerate(doc.tables):
        t_info = {
            'index': t_idx,
            'rows': len(table.rows),
            'cols': len(table.columns),
            'data': []
        }
        for r_idx, row in enumerate(table.rows):
            row_data = []
            for c_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()[:120]
                highlights = set()
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.highlight_color is not None:
                            highlights.add(str(run.font.highlight_color))
                row_data.append({
                    'text': cell_text,
                    'highlights': list(highlights) if highlights else [],
                })
            t_info['data'].append(row_data)
        result['tables'].append(t_info)

    return result


if __name__ == '__main__':
    print('Converting .doc to .docx...')
    convert_doc_to_docx(TEMPLATE_SRC, TEMPLATE_DOCX)
    print('Done.')

    print('Analyzing template...')
    result = analyze_template(TEMPLATE_DOCX)

    print(f"\nSections: {result['sections']}")

    print('\n=== Paragraphs with highlights ===')
    for p in result['paragraphs']:
        if p['highlights']:
            print(f"  [{p['index']}] {p['highlights']} | {p['text'][:150]}")

    print('\n=== Tables ===')
    for t in result['tables']:
        print(f"\nTable {t['index']}: {t['rows']}x{t['cols']}")
        for r_idx, row in enumerate(t['data']):
            has_hl = any(c['highlights'] for c in row)
            marker = ' *' if has_hl else ''
            texts = [c['text'][:40] for c in row]
            print(f"  R{r_idx}{marker}: {' | '.join(texts)}")

    # Save full analysis
    output_path = os.path.join(os.path.dirname(docx_path), 'template_analysis.json')
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    print(f'\nFull analysis saved to {output_path}')
