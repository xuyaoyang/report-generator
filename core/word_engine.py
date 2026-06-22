"""
Fill the prepared Word template with data from Excel.
Strategy: Use param_mapping.json location info to correctly map each
placeholder to its data source.
"""
import os
import re
import json
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _load_mapping(product_dir):
    """Load param_mapping.json from the product directory."""
    mapping_path = os.path.join(product_dir, 'param_mapping.json')
    with open(mapping_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def _parse_location(loc_str):
    """Parse a location string like 'table[6].cell[1,2]' or 'paragraph[10]'."""
    table_match = re.search(r'table\[(\d+)\]\.cell\[(\d+),(\d+)\]', loc_str)
    if table_match:
        return {
            'type': 'table',
            'table': int(table_match.group(1)),
            'row': int(table_match.group(2)),
            'col': int(table_match.group(3)),
        }
    para_match = re.search(r'paragraph\[(\d+)\]', loc_str)
    if para_match:
        return {
            'type': 'paragraph',
            'index': int(para_match.group(1)),
        }
    return {'type': 'unknown'}


def _normalize_report_month(report_date):
    """Normalize cover report date to YYYY-MM for output folder names."""
    text = str(report_date or '').strip()
    if not text:
        return '未知日期'

    match = re.search(r'(\d{4})\D*(\d{1,2})', text)
    if match:
        year = int(match.group(1))
        month = int(match.group(2))
        if 1 <= month <= 12:
            return f'{year:04d}-{month:02d}'

    compact = re.search(r'(\d{4})(\d{2})', text)
    if compact:
        year = int(compact.group(1))
        month = int(compact.group(2))
        if 1 <= month <= 12:
            return f'{year:04d}-{month:02d}'

    return (
        text.replace('年', '-')
        .replace('月', '')
        .replace('日', '')
        .replace('/', '-')
        .replace(' ', '')
        .rstrip('-')
    )


FP_KEYS = {
    'project': '\u9879\u76ee\u540d\u79f0',
    'supplier': '\u4ea7\u54c1\u4f9b\u5e94\u5546',
    'address': '\u5236\u9020\u5730\u5740',
    'phone': '\u8054\u7cfb\u7535\u8bdd',
    'date': '\u5236\u9020\u65e5\u671f',
    'inspector': '\u68c0\u9a8c\u5458',
    'model': '\u4ea7\u54c1\u578b\u53f7',
    'serial_range': '\u652f\u5ea7\u7f16\u53f7\u8303\u56f4',
    'qty': '\u6570\u91cf',
    'production_date': '\u751f\u4ea7\u65e5\u671f',
    'standard': '\u68c0\u9a8c\u6807\u51c6',
    'vertical': '\u7ad6\u5411\u627f\u8f7d\u529b(KN)',
    'displacement': '\u6c34\u5e73\u6781\u9650\u4f4d\u79fb(mm)',
    'period': '\u6446\u52a8\u5468\u671f(s)',
    'slow': '\u6469\u64e6\u7cfb\u6570(\u6162)',
    'fast': '\u6469\u64e6\u7cfb\u6570(\u5feb)',
    'height': '\u652f\u5ea7\u9ad8\u5ea6(mm)',
}


ED_KEYS = {
    'project': '\u9879\u76ee\u540d\u79f0',
    'supplier': '\u4ea7\u54c1\u4f9b\u5e94\u5546',
    'address': '\u5236\u9020\u5730\u5740',
    'phone': '\u8054\u7cfb\u7535\u8bdd',
    'date': '\u5236\u9020\u65e5\u671f',
    'embedded_type': '\u9884\u57cb\u4ef6\u7c7b\u578b',
    'inspector': '\u68c0\u9a8c\u5458',
    'reviewer': '\u5ba1\u6838',
    'issuer': '\u7b7e\u53d1',
    'product_name': '\u4ea7\u54c1\u540d\u79f0',
    'model': '\u4ea7\u54c1\u578b\u53f7',
    'qty': '\u6570\u91cf',
    'production_date': '\u751f\u4ea7\u65e5\u671f',
    'length': '\u957f\u5ea6\u5b9e\u6d4b\u503c',
    'width': '\u5bbd\u5ea6\u5b9e\u6d4b\u503c',
    'thickness': '\u539a\u5ea6\u5b9e\u6d4b\u503c',
}


IBE_KEYS = {
    'project': '\u9879\u76ee\u540d\u79f0',
    'supplier': '\u4ea7\u54c1\u4f9b\u5e94\u5546',
    'address': '\u5236\u9020\u5730\u5740',
    'phone': '\u8054\u7cfb\u7535\u8bdd',
    'date': '\u5236\u9020\u65e5\u671f',
    'inspector': '\u68c0\u9a8c\u5458',
    'reviewer': '\u5ba1\u6838',
    'issuer': '\u7b7e\u53d1',
    'bearing_model': '\u652f\u5ea7\u89c4\u683c',
    'bearing_qty': '\u652f\u5ea7\u6570\u91cf',
    'batch': '\u751f\u4ea7\u6279\u53f7',
    'production_date': '\u751f\u4ea7\u65e5\u671f',
}


def _parse_isolation_bearing_diameter(model):
    match = re.search(r'(?<!\d)(1[0-6]00|[4-9]00)(?!\d)', str(model or ''))
    if not match:
        raise ValueError(
            f'\u65e0\u6cd5\u4ece\u652f\u5ea7\u89c4\u683c“{model}”\u8bc6\u522b400~1600\u7684\u516c\u79f0\u76f4\u5f84')
    return match.group(1)


def _parse_whole_number(value, field_name):
    try:
        number = float(value)
    except (TypeError, ValueError):
        raise ValueError(f'{field_name}\u5fc5\u987b\u662f\u6570\u5b57')
    if number < 0 or not number.is_integer():
        raise ValueError(f'{field_name}\u5fc5\u987b\u662f\u975e\u8d1f\u6574\u6570')
    return int(number)


def _prepare_isolation_embedded_data(data, product_dir, material_manager=None):
    with open(os.path.join(product_dir, 'standard_connections.json'),
              'r', encoding='utf-8') as stream:
        standards = json.load(stream)

    project = data.get('project_info', {})
    default_date = str(project.get(IBE_KEYS['date'], ''))
    anchor_groups = {}
    plate_groups = {}

    for product in data.get('product_list', []):
        model = str(product.get(IBE_KEYS['bearing_model'], '')).strip()
        diameter = _parse_isolation_bearing_diameter(model)
        standard = standards.get(diameter)
        if not standard:
            raise ValueError(f'\u672a\u627e\u5230D{diameter}\u7684II\u578b\u8fde\u63a5\u53c2\u6570')
        qty = _parse_whole_number(
            product.get(IBE_KEYS['bearing_qty'], 0),
            f'{model}\u7684\u652f\u5ea7\u6570\u91cf')
        batch = str(product.get(IBE_KEYS['batch'], '')).strip()
        production_date = str(
            product.get(IBE_KEYS['production_date'], '') or default_date).strip()

        anchor_spec = (
            f'\u03c6{standard["anchor_diameter"]}\u00d7{standard["anchor_length"]}')
        anchor_key = (anchor_spec, batch, production_date)
        if anchor_key not in anchor_groups:
            anchor_groups[anchor_key] = {
                'name': '\u9884\u57cb\u4ef6',
                'spec': anchor_spec,
                'batch': batch,
                'production_date': production_date,
                'qty': 0,
                'diameter': standard['anchor_diameter'],
                'length': standard['anchor_length'],
                'sleeve_od': standard['sleeve_od'],
                'sleeve_length': standard['sleeve_length'],
                'material_grade': 'HRB400E',
                'manufacturer': '',
                'material_batch': '',
            }
        anchor_groups[anchor_key]['qty'] += qty * standard['multiplier']

        plate_spec = (
            f'{standard["plate_size"]}\u00d7{standard["plate_size"]}'
            f'\u00d7{standard["plate_thickness"]}')
        plate_key = (plate_spec, batch, production_date)
        if plate_key not in plate_groups:
            plate_groups[plate_key] = {
                'name': '\u9884\u57cb\u94a2\u677f',
                'spec': plate_spec,
                'batch': batch,
                'production_date': production_date,
                'qty': 0,
                'side': standard['plate_size'],
                'thickness': standard['plate_thickness'],
                'material_grade': 'Q235B',
                'manufacturer': '',
                'material_batch': '',
            }
        plate_groups[plate_key]['qty'] += qty

    anchors = list(anchor_groups.values())
    plates = list(plate_groups.values())
    cert_ids = []
    if material_manager:
        cert_cache = {}
        for item in anchors:
            key = ('\u94a2\u7b4b', item['diameter'], item['material_grade'])
            if key not in cert_cache:
                cert_cache[key] = material_manager.find_latest_certificate(
                    '\u94a2\u7b4b', {
                        'diameter': str(item['diameter']),
                        'material_grade': item['material_grade'],
                    })
            cert = cert_cache[key]
            if cert:
                item['material_grade'] = cert.get('params', {}).get(
                    'material_grade', item['material_grade'])
                item['manufacturer'] = cert.get('params', {}).get('manufacturer', '')
                item['material_batch'] = cert.get('batch_number', '')
                cert_ids.append(cert['id'])
        for item in plates:
            key = ('\u94a2\u677f', item['thickness'], item['material_grade'])
            if key not in cert_cache:
                cert_cache[key] = material_manager.find_latest_certificate(
                    '\u94a2\u677f', {
                        'thickness': str(item['thickness']),
                        'material_grade': item['material_grade'],
                    })
            cert = cert_cache[key]
            if cert:
                item['material_grade'] = cert.get('params', {}).get(
                    'material_grade', item['material_grade'])
                item['manufacturer'] = cert.get('params', {}).get('manufacturer', '')
                item['material_batch'] = cert.get('batch_number', '')
                cert_ids.append(cert['id'])

    prepared = dict(data)
    prepared['embedded_anchor_components'] = anchors
    prepared['embedded_plate_components'] = plates
    prepared['embedded_certificate_rows'] = anchors + plates
    prepared['_auto_material_cert_ids'] = list(dict.fromkeys(cert_ids))
    return prepared


def _build_isolation_embedded_common_fill_map(data):
    project = data.get('project_info', {})
    return {
        '{{IBE_TITLE}}': '\u9884\u57cb\u4ef6\u51fa\u5382\u8d44\u6599',
        '{{IBE_PROJECT}}': str(project.get(IBE_KEYS['project'], '')),
        '{{IBE_SUPPLIER}}': str(project.get(IBE_KEYS['supplier'], '')),
        '{{IBE_ADDRESS}}': str(project.get(IBE_KEYS['address'], '')),
        '{{IBE_PHONE}}': str(project.get(IBE_KEYS['phone'], '')),
        '{{IBE_DATE}}': str(project.get(IBE_KEYS['date'], '')),
        '{{IBE_INSPECTOR}}': str(project.get(IBE_KEYS['inspector'], '')),
        '{{IBE_REVIEWER}}': str(project.get(IBE_KEYS['reviewer'], '')),
        '{{IBE_ISSUER}}': str(project.get(IBE_KEYS['issuer'], '')),
    }


def _build_isolation_embedded_cert_fill_map(data, rows):
    fill = _build_isolation_embedded_common_fill_map(data)
    for index in range(1, 14):
        item = rows[index - 1] if index - 1 < len(rows) else {}
        fill.update({
            f'{{{{IBE_CERT_{index}_NAME}}}}': str(item.get('name', '/')),
            f'{{{{IBE_CERT_{index}_SPEC}}}}': str(item.get('spec', '/')),
            f'{{{{IBE_CERT_{index}_BATCH}}}}': str(item.get('batch', '/')),
            f'{{{{IBE_CERT_{index}_DATE}}}}': str(item.get('production_date', '/')),
            f'{{{{IBE_CERT_{index}_QTY}}}}': str(item.get('qty', '/')),
        })
    return fill


def _build_isolation_embedded_anchor_fill_map(data, item):
    fill = _build_isolation_embedded_common_fill_map(data)
    fill.update({
        '{{IBE_ANCHOR_SPEC}}': str(item.get('spec', '')),
        '{{IBE_ANCHOR_DIAMETER}}': str(item.get('diameter', '')),
        '{{IBE_ANCHOR_LENGTH}}': str(item.get('length', '')),
        '{{IBE_SLEEVE_OD}}': str(item.get('sleeve_od', '')),
        '{{IBE_SLEEVE_LENGTH}}': str(item.get('sleeve_length', '')),
        '{{IBE_ANCHOR_GRADE}}': str(item.get('material_grade', '')),
        '{{IBE_ANCHOR_MANUFACTURER}}': str(item.get('manufacturer', '')),
        '{{IBE_ANCHOR_BATCH}}': str(item.get('material_batch', '')),
    })
    return fill


def _build_isolation_embedded_plate_fill_map(data, item):
    fill = _build_isolation_embedded_common_fill_map(data)
    fill.update({
        '{{IBE_PLATE_SPEC}}': str(item.get('spec', '')),
        '{{IBE_PLATE_SIDE_STANDARD}}': f'{item.get("side", "")}\u00b12',
        '{{IBE_PLATE_THICKNESS_STANDARD}}': f'{item.get("thickness", "")}\u00b10.5',
        '{{IBE_PLATE_GRADE}}': str(item.get('material_grade', '')),
        '{{IBE_PLATE_MANUFACTURER}}': str(item.get('manufacturer', '')),
        '{{IBE_PLATE_BATCH}}': str(item.get('material_batch', '')),
    })
    return fill


def _build_embedded_damper_parts_fill_map(data):
    project = data.get('project_info', {})
    products = data.get('product_list', [])
    visual = data.get('visual_data', [])
    product = products[0] if products else {}
    model = str(product.get(ED_KEYS['model'], ''))
    vis_by_model = {str(v.get(ED_KEYS['model'], '')): v for v in visual}
    vis = vis_by_model.get(model, {})

    return {
        '{{ED_PROJECT}}': str(project.get(ED_KEYS['project'], '')),
        '{{ED_SUPPLIER}}': str(project.get(ED_KEYS['supplier'], '')),
        '{{ED_ADDRESS}}': str(project.get(ED_KEYS['address'], '')),
        '{{ED_PHONE}}': str(project.get(ED_KEYS['phone'], '')),
        '{{ED_DATE}}': str(project.get(ED_KEYS['date'], '')),
        '{{ED_BASE_PRODUCT}}': str(project.get(ED_KEYS['embedded_type'], '\u7c98\u6ede\u963b\u5c3c\u5668')),
        '{{ED_INSPECTOR}}': str(project.get(ED_KEYS['inspector'], '')),
        '{{ED_REVIEWER}}': str(project.get(ED_KEYS['reviewer'], '')),
        '{{ED_ISSUER}}': str(project.get(ED_KEYS['issuer'], '')),
        '{{ED_PRODUCT_NAME}}': str(product.get(ED_KEYS['product_name'], '\u9884\u57cb\u4ef6')),
        '{{ED_MODEL}}': model,
        '{{ED_QTY}}': str(product.get(ED_KEYS['qty'], '')),
        '{{ED_PRODUCTION_DATE}}': str(product.get(ED_KEYS['production_date'], '')),
        '{{ED_LENGTH}}': str(vis.get(ED_KEYS['length'], '')),
        '{{ED_WIDTH}}': str(vis.get(ED_KEYS['width'], '')),
        '{{ED_THICKNESS}}': str(vis.get(ED_KEYS['thickness'], '')),
    }


def _friction_product_name():
    return '\u5efa\u7b51\u6469\u64e6\u6446\u9694\u9707\u652f\u5ea7'


def _build_friction_fill_map(data):
    project = data.get('project_info', {})
    products = data.get('product_list', [])
    mechanical = data.get('mechanical_data', [])
    mech_by_model = {str(m.get(FP_KEYS['model'], '')): m for m in mechanical}
    product = products[0] if products else {}

    fill = {
        '{{FP_PROJECT}}': str(project.get(FP_KEYS['project'], '')),
        '{{FP_SUPPLIER}}': str(project.get(FP_KEYS['supplier'], '')),
        '{{FP_ADDRESS}}': str(project.get(FP_KEYS['address'], '')),
        '{{FP_PHONE}}': str(project.get(FP_KEYS['phone'], '')),
        '{{FP_DATE}}': str(project.get(FP_KEYS['date'], '')),
        '{{FP_MODEL}}': str(product.get(FP_KEYS['model'], '')),
        '{{FP_SERIAL_RANGE}}': str(product.get(FP_KEYS['serial_range'], '')),
        '{{FP_PRODUCTION_DATE}}': str(product.get(FP_KEYS['production_date'], '')),
        '{{FP_STANDARD}}': str(product.get(FP_KEYS['standard'], 'GB/T37358-2019')),
        '{{FP_INSPECTOR}}': str(project.get(FP_KEYS['inspector'], '')),
    }

    qty = product.get(FP_KEYS['qty'], '')
    fill['{{FP_QUANTITY_WITH_UNIT}}'] = f'{qty}\u4ef6' if str(qty) and '\u4ef6' not in str(qty) else str(qty)
    total_qty = sum(int(_parse_numeric(p.get(FP_KEYS['qty']), 0)) for p in products)
    fill['{{FP_TOTAL_QTY}}'] = f'{total_qty}\u5957' if total_qty else ''

    for idx in range(1, 7):
        p = products[idx - 1] if idx - 1 < len(products) else {}
        model = str(p.get(FP_KEYS['model'], ''))
        mech = mech_by_model.get(model, {})
        fill[f'{{{{FP_SUM_{idx}_MODEL}}}}'] = model
        fill[f'{{{{FP_SUM_{idx}_VERTICAL}}}}'] = str(mech.get(FP_KEYS['vertical'], ''))
        fill[f'{{{{FP_SUM_{idx}_DISP}}}}'] = str(mech.get(FP_KEYS['displacement'], ''))
        fill[f'{{{{FP_SUM_{idx}_PERIOD}}}}'] = str(mech.get(FP_KEYS['period'], ''))
        fill[f'{{{{FP_SUM_{idx}_SLOW}}}}'] = str(mech.get(FP_KEYS['slow'], ''))
        fill[f'{{{{FP_SUM_{idx}_FAST}}}}'] = str(mech.get(FP_KEYS['fast'], ''))
        fill[f'{{{{FP_SUM_{idx}_HEIGHT}}}}'] = str(mech.get(FP_KEYS['height'], ''))
        fill[f'{{{{FP_SUM_{idx}_QTY}}}}'] = str(mech.get(FP_KEYS['qty'], p.get(FP_KEYS['qty'], '')))

    for idx in range(1, 17):
        fill[f'{{{{FP_DET_{idx}_MODEL}}}}'] = ''
        fill[f'{{{{FP_DET_{idx}_SERIAL}}}}'] = ''
        fill[f'{{{{FP_DET_{idx}_RESULT}}}}'] = ''
        fill[f'{{{{FP_DET_{idx}_SLASH}}}}'] = ''

    return fill


def _expand_friction_serials(products):
    rows = []
    for product in products:
        model = str(product.get(FP_KEYS['model'], ''))
        serial_range = str(product.get(FP_KEYS['serial_range'], ''))
        match = re.match(r'([A-Za-z]+)(\d+)\s*-\s*([A-Za-z]+)?(\d+)', serial_range)
        if match:
            prefix1, start, prefix2, end = match.groups()
            prefix = prefix1
            start_i = int(start)
            end_i = int(end)
            width = len(start)
            for num in range(start_i, end_i + 1):
                rows.append({'model': model, 'serial': f'{prefix}{num:0{width}d}'})
            continue
        qty = int(_parse_numeric(product.get(FP_KEYS['qty']), 0))
        count = max(qty, 1)
        for _ in range(count):
            rows.append({'model': model, 'serial': serial_range})
    return rows


def _build_friction_detail_fill_map(rows):
    fill = {}
    for idx in range(1, 17):
        row = rows[idx - 1] if idx - 1 < len(rows) else {}
        has_row = bool(row.get('model') or row.get('serial'))
        fill[f'{{{{FP_DET_{idx}_MODEL}}}}'] = row.get('model', '')
        fill[f'{{{{FP_DET_{idx}_SERIAL}}}}'] = row.get('serial', '')
        fill[f'{{{{FP_DET_{idx}_RESULT}}}}'] = '\u5408\u683c' if has_row else ''
        fill[f'{{{{FP_DET_{idx}_SLASH}}}}'] = '/' if has_row else ''
    return fill


def build_fill_map(data, product_dir=None, product_type='isolation_bearing'):
    """
    Build a mapping from placeholder name to replacement value.
    Uses param_mapping.json to correctly identify what each
    placeholder represents based on its location in the template.

"""
    if product_type == 'friction_pendulum':
        return _build_friction_fill_map(data)
    if product_type == 'embedded_damper_parts':
        return _build_embedded_damper_parts_fill_map(data)
    if product_type == 'isolation_bearing_embedded_parts':
        return _build_isolation_embedded_common_fill_map(data)

    if product_dir is None:
        product_dir = os.path.join(
            os.path.dirname(os.path.dirname(__file__)),
            'products', 'isolation_bearing'
        )

    mapping = _load_mapping(product_dir)
    fill = {}

    project = data.get('project_info', {})
    products = data.get('product_list', [])
    mechanical = data.get('mechanical_data', [])
    visual = data.get('visual_data', [])

    mech_by_model = {m['产品型号']: m for m in mechanical}
    vis_by_model = {v['产品型号']: v for v in visual}

    # ---- Cover info (paragraph fields) ----
    # Label → project_info key mapping (covers both isolation bearing and viscous damper)
    COVER_LABEL_TO_KEY = {
        '项目名称': '项目名称',
        '产品供应商': '产品供应商',
        '生产地址': '制造地址',
        '制造地址': '制造地址',
        '电话': '联系电话',
        '联系电话': '联系电话',
        '报告日期': '制造日期',
        '制造日期': '制造日期',
    }

    for f in mapping['categories']['cover_info']:
        loc_str = f['location']
        loc = _parse_location(loc_str)
        p_idx = loc.get('index', -1)

        # Try label-based matching first (for templates with labels in parens)
        label_match = re.search(r'\(([^)]+)\)', loc_str)
        if label_match:
            label = label_match.group(1)
            key = COVER_LABEL_TO_KEY.get(label, label)
            fill[f['placeholder']] = str(project.get(key, ''))
        elif p_idx == 10:
            fill[f['placeholder']] = str(project.get('项目名称', ''))
        elif p_idx == 14:
            fill[f['placeholder']] = str(project.get('产品供应商', ''))
        elif p_idx == 15:
            fill[f['placeholder']] = str(project.get('制造地址', ''))
        elif p_idx == 17:
            fill[f['placeholder']] = str(project.get('制造日期', ''))
        elif p_idx == 101:
            if 'FIELD_145' in f['placeholder']:
                fill[f['placeholder']] = str(project.get('检验员', ''))
            elif 'FIELD_146' in f['placeholder']:
                fill[f['placeholder']] = str(project.get('审核', ''))
            else:
                fill[f['placeholder']] = ''
        else:
            fill[f['placeholder']] = ''

    # ---- Certificates ----
    # Isolation bearing: each field is on a different row (row-only mapping).
    # Viscous damper: fields are at specific (row, col) positions (row+col mapping).
    IB_CERT_ROW_MAP = {1: '产品型号', 2: '支座编号范围', 3: '生产日期', 4: '数量', 5: '检验标准'}
    # VD cert uses both row+col; detect by col being 1 or 3
    for f in mapping['categories']['certificates']:
        loc = _parse_location(f['location'])
        table_idx = loc['table']
        row = loc['row']
        col = loc.get('col', 0)

        if product_type == 'viscous_damper' and col in (1, 3) and row in (0, 1, 2, 3):
            # VD cert format: (row, col) → field
            VD_CERT_MAP = {
                (0, 3): ('产品型号', 'products'),
                (1, 1): ('生产日期', 'products'),
                (1, 3): ('阻尼器编号范围', 'products'),
                (2, 1): ('检验员', 'project'),
                (2, 3): ('检验标准', 'products'),
                (3, 1): ('检验部门', 'static'),
                (3, 3): ('检验结果', 'static'),
            }
            key, source = VD_CERT_MAP.get((row, col), ('', ''))
            if source == 'products' and table_idx < len(products):
                fill[f['placeholder']] = str(products[table_idx].get(key, ''))
            elif source == 'project':
                fill[f['placeholder']] = str(project.get(key) or '')
            elif source == 'static':
                fill[f['placeholder']] = '质量部' if key == '检验部门' else '合格'
            else:
                fill[f['placeholder']] = ''
        else:
            # IB cert format: row-only mapping
            cert_field = IB_CERT_ROW_MAP.get(row)
            if cert_field and table_idx < len(products):
                fill[f['placeholder']] = str(products[table_idx].get(cert_field, ''))
            else:
                fill[f['placeholder']] = ''

    # ---- Mechanical summary (Table 6) ----
    # Table 6 structure:
    #   Row 1, cols 2-12: project name (merged)
    #   Row 3, cols 2-12: supplier (merged)
    #   Row 4, cols 10-12: total quantity (merged)
    #   Rows 6-11, cols 1-12: mechanical design values for 6 models
    #     col1=型号, col2=有效直径, col3=高度, col4=橡胶层厚, col5=压缩应力,
    #     col6=设计压缩刚度, col7-9=水平等效刚度, col10=屈服后刚度, col11=屈服力, col12=阻尼比
    MECH_COL_MAP = {
        1: '产品型号',
        2: '支座有效直径(mm)',
        3: '支座高度(mm)',
        4: '橡胶层总厚(mm)',
        5: '压缩应力(MPa)',
        6: '设计压缩刚度(KN/mm)',
        7: '水平等效刚度(KN/mm)',
        8: '水平等效刚度(KN/mm)',
        9: '水平等效刚度(KN/mm)',
        10: '屈服后刚度(KN/mm)',
        11: '屈服力(KN)',
        12: '等效阻尼比(%)',
    }

    for f in mapping['categories']['mechanical']:
        loc = _parse_location(f['location'])
        if loc['type'] != 'table':
            continue
        t = loc['table']
        r = loc['row']
        c = loc['col']

        if t == 6:
            if r == 1:
                # Project name row (cols 2-12)
                fill[f['placeholder']] = str(project.get('项目名称', ''))
            elif r == 3:
                # Supplier row (cols 2-12)
                fill[f['placeholder']] = str(project.get('产品供应商', ''))
            elif r == 4:
                # Quantity row (cols 10-12) — use total count or model-specific
                fill[f['placeholder']] = str(project.get('数量', sum(
                    int(p.get('数量', 0)) for p in products if str(p.get('数量', '')).isdigit()
                )))
            elif 6 <= r <= 11:
                # Data rows: r6=model0, r7=model1, ..., r11=model5
                model_idx = r - 6
                col_key = MECH_COL_MAP.get(c)
                if col_key and model_idx < len(products):
                    model_name = products[model_idx]['产品型号']
                    m = mech_by_model.get(model_name, {})
                    val = m.get(col_key, '/')
                    fill[f['placeholder']] = str(val) if val is not None else '/'
                else:
                    fill[f['placeholder']] = ''
            else:
                fill[f['placeholder']] = ''

        elif t in (7, 8):
            # Mechanical detail tables
            # Data rows at 3,7,11,15 (every 4th row starting at 3)
            model_idx_map = {3: 0, 7: 1, 11: 2, 15: 3}
            model_idx = model_idx_map.get(r, -1)

            # Table 7 → models 0-3, Table 8 → models 4-7
            mech_page_offset = 0 if t == 7 else 4
            global_model_idx = model_idx + mech_page_offset

            if global_model_idx < len(products):
                p = products[global_model_idx]
                model_name = p.get('产品型号', '')
                mech = mech_by_model.get(model_name, {})

                if c == 0:
                    fill[f['placeholder']] = str(p.get('产品型号', ''))
                elif c == 1:
                    fill[f['placeholder']] = str(p.get('支座编号范围', ''))
                elif c >= 2:
                    # Judgment columns
                    judgment_keys = [
                        '设计压缩刚度(KN/mm)', '水平等效刚度(KN/mm)',
                        '屈服后刚度(KN/mm)', '屈服力(KN)', '等效阻尼比(%)'
                    ]
                    if c - 2 < len(judgment_keys):
                        key = judgment_keys[c - 2]
                        val = mech.get(key, '/') if mech else '/'
                        fill[f['placeholder']] = _compute_judgment(val)
                    else:
                        fill[f['placeholder']] = '/'
            else:
                fill[f['placeholder']] = ''

        elif t == 9:
            VD_MECH_DESIGN_KEYS = {
                0: '最大阻尼力(kN)',
                1: '设计位移(mm)',
                2: '极限位移(mm)',
                3: '阻尼系数(kN(s/mm)α)',
                4: '阻尼力指数(α)',
                5: '结构基频(Hz)',
            }
            VD_MECH_RESULT_ROWS = {7: '设计阻尼力', 8: '极限位移', 9: '阻尼系数', 10: '阻尼指数'}
            first_product = products[0] if products else {}
            first_mech = mech_by_model.get(first_product.get('产品型号', ''), {})

            if r == 0:
                fill[f['placeholder']] = str(project.get('项目名称', ''))
            elif r == 1:
                fill[f['placeholder']] = str(first_product.get('产品型号', ''))
            elif r == 4:
                key = VD_MECH_DESIGN_KEYS.get(c)
                if key:
                    val = first_mech.get(key, '') if first_mech else ''
                    fill[f['placeholder']] = str(val) if val else ''
                else:
                    fill[f['placeholder']] = ''
            elif r in VD_MECH_RESULT_ROWS:
                # Merged cells: only col 4 placeholders exist in prepared template.
                # The merged cell covers both 检测结果 and 判定 columns.
                label = VD_MECH_RESULT_ROWS[r]
                key = f'{label}(合格/不合格)'
                verdict = str(first_mech.get(key, '合格'))
                fill[f['placeholder']] = '符合' if verdict == '合格' else '不符合'
            else:
                fill[f['placeholder']] = ''

        elif 9 <= t <= 14:
            # IB visual inspection tables
            # Table 9=model0, Table10=model1, ..., Table14=model5
            model_idx = t - 9
            if r == 0:
                # Project name
                fill[f['placeholder']] = str(project.get('项目名称', ''))
            elif r == 1:
                # Quantity for this model
                if model_idx < len(products):
                    fill[f['placeholder']] = str(products[model_idx].get('数量', ''))
                else:
                    fill[f['placeholder']] = ''
            else:
                fill[f['placeholder']] = ''
        else:
            fill[f['placeholder']] = ''

    # ---- Visual inspection (tables 9-14 and paragraph headers) ----
    VISUAL_PARA_LABEL_TO_KEY = {
        '外观报告-公司名称': '产品供应商',
        '外观报告-检验员': '检验员',
        '外观报告-审核': '审核',
        '外观报告-检测日期': '制造日期',
    }

    for f in mapping['categories']['visual']:
        loc = _parse_location(f['location'])
        if loc['type'] == 'paragraph':
            label_match = re.search(r'\(([^)]+)\)', f['location'])
            if label_match:
                label = label_match.group(1)
                key = VISUAL_PARA_LABEL_TO_KEY.get(label, label)
                fill[f['placeholder']] = str(project.get(key, ''))
            else:
                fill[f['placeholder']] = ''
            continue
        if loc['type'] != 'table':
            continue
        t = loc['table']

        if t == 8:
            # Viscous damper shared visual table — label-based filling
            label_match = re.search(r'\(([^)]+)\)', f['location'])
            label = label_match.group(1) if label_match else ''
            first_model = products[0] if products else {}
            first_vis = vis_by_model.get(first_model.get('产品型号', ''), {}) if products else {}
            model_name = str(first_model.get('产品型号', ''))
            conclusion_text = (
                f'检验结论：\n\n该项目黏滞阻尼器 {model_name} '
                '所检项目符合<<JG/T209-2012>>中质量标准要求，产品均检验合格 。'
            )

            VISUAL_ITEM_TO_KEY = {
                '表面平滑度': '表面平滑度(合格/不合格)',
                '机械损伤': '机械损伤(合格/不合格)',
                '锈蚀毛刺': '锈蚀毛刺(合格/不合格)',
                '无渗漏': '无渗漏(合格/不合格)',
                '产品标识': '产品标识(合格/不合格)',
                '长度偏差': '长度偏差(合格/不合格)',
                '截面尺寸偏差': '截面尺寸偏差(合格/不合格)',
            }

            if '项目名称' in label:
                fill[f['placeholder']] = str(project.get('项目名称', ''))
            elif '产品型号' in label or '检验结论-型号' in label:
                fill[f['placeholder']] = model_name
            elif '检验结论-项目名称' in label:
                fill[f['placeholder']] = '该项目黏滞阻尼器'
            elif '检验结论' in label:
                fill[f['placeholder']] = conclusion_text
            elif '检验部门' in label:
                fill[f['placeholder']] = conclusion_text if loc.get('row') == 11 else '质量部'
            elif '公司名称' in label:
                fill[f['placeholder']] = str(project.get('产品供应商', ''))
            elif '检验员' in label:
                fill[f['placeholder']] = str(project.get('检验员', ''))
            elif '审核' in label:
                fill[f['placeholder']] = str(project.get('审核', ''))
            elif '检测日期' in label:
                fill[f['placeholder']] = str(project.get('制造日期', ''))
            elif '数量' in label:
                fill[f['placeholder']] = str(first_model.get('数量', ''))
            elif '实测值' in label:
                # Derive descriptive measured value from verdict
                VD_MEASURED_MAP = {
                    '表面平滑度': {'合格': '光滑平整', '不合格': '不平整'},
                    '机械损伤': {'合格': '无', '不合格': '有'},
                    '锈蚀毛刺': {'合格': '无', '不合格': '有'},
                    '无渗漏': {'合格': '无', '不合格': '有渗漏'},
                    '产品标识': {'合格': '标识清晰易辨认', '不合格': '标识不清晰'},
                    '长度偏差': {'合格': '符合标准要求', '不合格': '超出允许偏差'},
                    '截面尺寸偏差': {'合格': '符合标准要求', '不合格': '超出允许偏差'},
                }
                for item_name, data_key in VISUAL_ITEM_TO_KEY.items():
                    if item_name in label:
                        verdict = str(first_vis.get(data_key, '合格'))
                        desc = VD_MEASURED_MAP.get(item_name, {})
                        fill[f['placeholder']] = desc.get(verdict, verdict)
                        break
                else:
                    fill[f['placeholder']] = '符合'
            elif '结论' in label:
                for item_name, data_key in VISUAL_ITEM_TO_KEY.items():
                    if item_name in label:
                        val = first_vis.get(data_key, '合格')
                        fill[f['placeholder']] = str(val) if val else '合格'
                        break
                else:
                    fill[f['placeholder']] = '合格'
            else:
                fill[f['placeholder']] = ''
        else:
            r = loc['row']
            model_idx = t - 9  # Table 9 → model 0, Table 14 → model 5

            if r == 0:
                fill[f['placeholder']] = str(project.get('项目名称', ''))
            elif r == 1:
                if model_idx < len(products):
                    fill[f['placeholder']] = str(products[model_idx].get('数量', ''))
                else:
                    fill[f['placeholder']] = ''

    # ---- Ensure all 202 fields have at least an empty string ----
    for i in range(1, 203):
        key = f'{{{{FIELD_{i:03d}}}}}'
        if key not in fill:
            fill[key] = ''

    return fill


def fill_template(template_path, output_path, fill_map):
    """
    Open the prepared template, replace all {{FIELD_XXX}} placeholders
    with values from fill_map, and save to output_path.
    """
    doc = Document(template_path)

    def replace_in_paragraphs(paragraphs):
        for para in paragraphs:
            for run in para.runs:
                for placeholder, value in fill_map.items():
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)

    # Replace in body paragraphs
    replace_in_paragraphs(doc.paragraphs)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs)

    # Replace in headers/footers
    for section in doc.sections:
        replace_in_paragraphs(section.header.paragraphs)
        replace_in_paragraphs(section.footer.paragraphs)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def _set_cell_text(cell_elem, new_text):
    """Directly set all text in an XML table cell element."""
    first = True
    for t in cell_elem.iter(qn('w:t')):
        if t.text is not None:
            if first:
                t.text = new_text
                first = False
            else:
                t.text = ''


def _compute_judgment(val):
    """Convert a numeric value to judgment text ('合格' or '/')."""
    if val is None:
        return '/'
    s = str(val).strip().replace('%', '')
    if s in ('/', '0', '', 'None', '.'):
        return '/'
    try:
        float(s)
        return '合格'
    except ValueError:
        return str(val)


def _set_keep_with_next(para_elem):
    """Add <w:keepNext/> to a paragraph's properties."""
    pPr = para_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para_elem.insert(0, pPr)
    if pPr.find(qn('w:keepNext')) is None:
        pPr.append(OxmlElement('w:keepNext'))


def _set_cant_split(tr_elem):
    """Add <w:cantSplit/> to a table row, preventing it from breaking across pages."""
    trPr = tr_elem.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr_elem.insert(0, trPr)
    if trPr.find(qn('w:cantSplit')) is None:
        trPr.append(OxmlElement('w:cantSplit'))


def _fix_certificate_pagination(doc):
    """Prevent certificate groups (title + table) from splitting across pages.
    Sets 'keep with next' on title paragraphs and 'cantSplit' on table rows."""
    body = doc.element.body
    body_children = list(body)

    for i, child in enumerate(body_children):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'p':
            continue
        text = ''
        for t in child.iter(qn('w:t')):
            if t.text:
                text += t.text
        if text.strip() != '产品合格证':
            continue

        # Keep title with the following content
        _set_keep_with_next(child)

        # Find the certificate table (skip empty paragraphs)
        tbl_elem = None
        for j in range(i + 1, min(len(body_children), i + 5)):
            nc = body_children[j]
            nc_tag = nc.tag.split('}')[-1] if '}' in nc.tag else nc.tag
            if nc_tag == 'p':
                nc_text = ''
                for t in nc.iter(qn('w:t')):
                    if t.text:
                        nc_text += t.text
                if not nc_text.strip():
                    _set_keep_with_next(nc)
                    continue
                break
            elif nc_tag == 'tbl':
                tbl_elem = nc
                break

        if tbl_elem is None:
            continue

        # Set cantSplit on all rows; set keepNext on every paragraph
        # inside every cell of all rows except the last.
        tr_elems = tbl_elem.findall(qn('w:tr'))
        for idx, tr in enumerate(tr_elems):
            _set_cant_split(tr)
            if idx < len(tr_elems) - 1:
                for tc in tr.findall(qn('w:tc')):
                    for p in tc.findall(qn('w:p')):
                        _set_keep_with_next(p)


def _direct_fill_all_tables(doc, data):
    """Post-processing: directly fill table cells that have no placeholders.
    Covers:
    1. Mechanical detail tables — all judgment cells (fixes hardcoded '/' in Table 7)
    2. Visual inspection tables — row 3 model name + last row '工程项目XXX'
    """
    products = data.get('product_list', [])
    mechanical = data.get('mechanical_data', [])
    mech_by_model = {m['产品型号']: m for m in mechanical}

    JUDGMENT_KEYS = [
        '设计压缩刚度(KN/mm)', '水平等效刚度(KN/mm)',
        '屈服后刚度(KN/mm)', '屈服力(KN)', '等效阻尼比(%)'
    ]
    MODELS_PER_PAGE = 4
    body = doc.element.body
    body_children = list(body)

    # ================================================================
    # 1. Mechanical detail tables (pattern: company → title → table)
    # ================================================================
    mech_page = 0
    for i, child in enumerate(body_children):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'p':
            continue
        text = ''
        for t in child.iter(qn('w:t')):
            if t.text:
                text += t.text
        if '四川融海运通' not in text:
            continue

        if i + 2 >= len(body_children):
            continue
        n1 = body_children[i + 1]
        if (n1.tag.split('}')[-1] if '}' in n1.tag else n1.tag) != 'p':
            continue
        n1_text = ''
        for t in n1.iter(qn('w:t')):
            if t.text:
                n1_text += t.text
        if '隔震支座力学性能检验报告' not in n1_text:
            continue

        n2 = body_children[i + 2]
        if (n2.tag.split('}')[-1] if '}' in n2.tag else n2.tag) != 'tbl':
            continue

        # Found a mechanical detail table
        trs = n2.findall(qn('w:tr'))
        if len(trs) > 16:
            continue  # skip summary table (17 rows)

        model_offset = mech_page * MODELS_PER_PAGE
        mech_page += 1

        for row_idx in [3, 7, 11, 15]:
            if row_idx >= len(trs):
                continue
            m_idx = model_offset + {3: 0, 7: 1, 11: 2, 15: 3}[row_idx]
            if m_idx >= len(products):
                continue

            product = products[m_idx]
            model_name = product.get('产品型号', '')
            mech = mech_by_model.get(model_name, {})

            cells = trs[row_idx].findall(qn('w:tc'))
            if len(cells) > 0:
                _set_cell_text(cells[0], str(model_name))
            if len(cells) > 1:
                _set_cell_text(cells[1], str(product.get('支座编号范围', '')))
            for col_idx in range(2, min(7, len(cells))):
                key = JUDGMENT_KEYS[col_idx - 2]
                val = mech.get(key, '/') if mech else '/'
                _set_cell_text(cells[col_idx], _compute_judgment(val))

    # ================================================================
    # 2. Visual inspection tables (pattern: title → ... → table)
    # ================================================================
    vis_page = 0
    for i, child in enumerate(body_children):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'p':
            continue
        text = ''
        for t in child.iter(qn('w:t')):
            if t.text:
                text += t.text
        if '隔震橡胶支座外观质量及尺寸检测报告' not in text:
            continue

        # Find the associated table
        tbl_elem = None
        for j in range(i + 1, min(len(body_children), i + 5)):
            nc = body_children[j]
            if (nc.tag.split('}')[-1] if '}' in nc.tag else nc.tag) == 'tbl':
                tbl_elem = nc
                break
        if tbl_elem is None:
            continue

        trs = tbl_elem.findall(qn('w:tr'))
        if len(trs) < 21:
            continue

        if vis_page >= len(products):
            vis_page += 1
            continue

        product = products[vis_page]
        model_name = product.get('产品型号', '')

        # Row 3, col 3: read OLD model name first, then set new
        cells_r3 = trs[3].findall(qn('w:tc'))
        old_model_name = ''
        if len(cells_r3) > 3:
            for t in cells_r3[3].iter(qn('w:t')):
                if t.text:
                    old_model_name += t.text
        old_model_name = old_model_name.strip()

        if len(cells_r3) > 3:
            _set_cell_text(cells_r3[3], model_name)

        # Last row: replace old model name with new, preserving all
        # other w:t runs (spacing, labels, etc.) untouched.
        if old_model_name:
            last_tr = trs[-1]
            for cell in last_tr.findall(qn('w:tc')):
                all_ts = list(cell.iter(qn('w:t')))
                texts = [t.text if t.text else '' for t in all_ts]
                combined = ''.join(texts)
                pos = combined.find(old_model_name)
                if pos < 0:
                    continue

                # Find which w:t elements cover the old model name
                running = 0
                start_elem = -1
                end_elem = -1
                for idx, txt in enumerate(texts):
                    if start_elem < 0 and running + len(txt) > pos:
                        start_elem = idx
                    if start_elem >= 0:
                        running += len(txt)
                        if running >= pos + len(old_model_name):
                            end_elem = idx
                            break
                    else:
                        running += len(txt)

                if start_elem >= 0 and end_elem >= start_elem:
                    all_ts[start_elem].text = model_name
                    for idx in range(start_elem + 1, end_elem + 1):
                        if all_ts[idx].text is not None:
                            all_ts[idx].text = ''

        vis_page += 1


def _elem_tag(elem):
    return elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag


def _elem_text(elem):
    """Get concatenated w:t text from an element."""
    text = ''
    for t in elem.iter(qn('w:t')):
        if t.text:
            text += t.text
    return text.strip()


def _has_page_break(elem):
    """Check if an element contains a hard page break."""
    return any(br.get(qn('w:type')) == 'page' for br in elem.iter(qn('w:br')))


def _has_image(elem):
    """Check if an element contains a drawing, VML picture, or embedded object."""
    return any(elem.findall('.//' + qn(tag)) for tag in ('w:drawing', 'w:pict', 'w:object'))


def _has_section_break(elem):
    """Check if an element marks the end of a document section."""
    pPr = elem.find(qn('w:pPr'))
    return pPr is not None and pPr.find(qn('w:sectPr')) is not None


def _fix_section_breaks(doc):
    """Remove redundant w:br page breaks from paragraphs that already have
    a section break (w:sectPr in w:pPr). The section break's NEW_PAGE type
    already provides the page break; having both creates a blank page."""
    body = doc.element.body
    removed = 0
    for child in list(body):
        if _elem_tag(child) != 'p':
            continue
        pPr = child.find(qn('w:pPr'))
        if pPr is None:
            continue
        sectPr = pPr.find(qn('w:sectPr'))
        if sectPr is None:
            continue
        # This paragraph has a section break — remove any w:br page breaks
        for run in child.findall(qn('w:r')):
            for br in run.findall(qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    run.remove(br)
                    removed += 1
    if removed:
        print(f'Removed {removed} redundant page break(s) from section-break paragraphs.')


def _normalize_page_orientation(doc):
    """Ensure all sections use portrait orientation.
    Some sections in the viscous damper template's mechanical report
    may be landscape; swap them to portrait."""
    fixed = 0
    for section in doc.sections:
        sectPr = section._sectPr
        pgSz = sectPr.find(qn('w:pgSz'))
        if pgSz is None:
            # Section has no explicit pgSz — add portrait A4
            pgSz = OxmlElement('w:pgSz')
            pgSz.set(qn('w:w'), '11906')
            pgSz.set(qn('w:h'), '16839')
            sectPr.append(pgSz)
            fixed += 1
            continue

        w_val = pgSz.get(qn('w:w'))
        h_val = pgSz.get(qn('w:h'))
        orient = pgSz.get(qn('w:orient'))

        if w_val and h_val:
            w = int(w_val)
            h = int(h_val)
            if w > h:
                pgSz.set(qn('w:w'), str(h))
                pgSz.set(qn('w:h'), str(w))
                orient_key = qn('w:orient')
                if orient_key in pgSz.attrib:
                    del pgSz.attrib[orient_key]
                fixed += 1
            elif orient == 'landscape':
                del pgSz.attrib[qn('w:orient')]
                fixed += 1

    if fixed:
        print(f'Normalized {fixed} section(s) from landscape to portrait.')


def _fit_tables_to_portrait(doc):
    """Resize tables that are wider than portrait A4 width (9072 dxa).
    Tables designed for landscape pages may have fixed widths that exceed
    portrait page width, causing content to be cut off."""
    PORTRAIT_DXA = 9072
    resized = 0
    for tbl in doc.element.body.iter(qn('w:tbl')):
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            continue
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            continue
        w_type = tblW.get(qn('w:type'), '')
        w_val = tblW.get(qn('w:w'), '')
        if w_type == 'dxa' and w_val:
            if int(w_val) > PORTRAIT_DXA:
                tblW.set(qn('w:w'), str(PORTRAIT_DXA))
                resized += 1
    if resized:
        print(f'Resized {resized} table(s) to fit portrait width.')


def _compact_certificates(doc):
    """Make certificate tables more compact so at least 2 fit per page.
    - Reduce row heights from 603 to 320 twips
    - Reduce paragraph spacing in cells
    - Remove empty paragraphs between certificate groups
    Preserves paragraphs that contain page breaks.
    """
    body = list(doc.element.body)

    # ---- Step 1: find cert groups (title_index, table_index) ----
    cert_groups = []
    for i, child in enumerate(body):
        if _elem_tag(child) != 'p':
            continue
        if _elem_text(child) != '产品合格证':
            continue
        tbl_idx = None
        for j in range(i + 1, min(len(body), i + 5)):
            if _elem_tag(body[j]) == 'tbl':
                tbl_idx = j
                break
        if tbl_idx is not None:
            cert_groups.append((i, tbl_idx))

    # ---- Step 2: compact each cert table ----
    for _, tbl_idx in cert_groups:
        tbl_elem = body[tbl_idx]
        for tr in tbl_elem.findall(qn('w:tr')):
            trPr = tr.find(qn('w:trPr'))
            if trPr is None:
                trPr = OxmlElement('w:trPr')
                tr.insert(0, trPr)
            trHeight = trPr.find(qn('w:trHeight'))
            if trHeight is None:
                trHeight = OxmlElement('w:trHeight')
                trPr.append(trHeight)
            trHeight.set(qn('w:val'), '360')
            trHeight.set(qn('w:hRule'), 'atLeast')

            for tc in tr.findall(qn('w:tc')):
                for p in tc.findall(qn('w:p')):
                    pPr = p.find(qn('w:pPr'))
                    if pPr is None:
                        pPr = OxmlElement('w:pPr')
                        p.insert(0, pPr)
                    spacing = pPr.find(qn('w:spacing'))
                    if spacing is None:
                        spacing = OxmlElement('w:spacing')
                        pPr.append(spacing)
                    spacing.set(qn('w:before'), '0')
                    spacing.set(qn('w:after'), '0')
                    spacing.set(qn('w:line'), '280')
                    spacing.set(qn('w:lineRule'), 'auto')

    # ---- Step 3: remove empty paragraphs between cert groups ----
    # Only remove empty paragraphs WITHOUT page breaks
    removed = set()
    for idx in range(len(cert_groups) - 1):
        _, cur_tbl = cert_groups[idx]
        next_title, _ = cert_groups[idx + 1]
        for k in range(cur_tbl + 1, next_title):
            child = body[k]
            if _elem_tag(child) == 'p' and not _elem_text(child) and not _has_page_break(child) and not _has_image(child) and not _has_section_break(child):
                removed.add(child)

    # Also remove empty paragraphs after last cert table (but before
    # the next table/section — stop at page breaks or non-empty content)
    if cert_groups:
        _, last_tbl = cert_groups[-1]
        for k in range(last_tbl + 1, len(body)):
            tag = _elem_tag(body[k])
            if tag == 'tbl':
                break
            if tag == 'p':
                text = _elem_text(body[k])
                if not text and not _has_page_break(body[k]) and not _has_image(body[k]) and not _has_section_break(body[k]):
                    removed.add(body[k])
                else:
                    break  # stop at page break or non-empty content

    for elem in removed:
        doc.element.body.remove(elem)

    print(f'Compacted certificates: reduced row heights, removed '
          f'{len(removed)} empty paragraphs.')


def _remove_blank_pages(doc):
    """Collapse runs of consecutive empty paragraphs to avoid blank pages.
    Keeps at most 1 empty paragraph in a row. Preserves page breaks."""
    body_children = list(doc.element.body)

    # Find runs of 2+ consecutive empty paragraphs
    runs_to_collapse = []
    run_start = None
    for i, child in enumerate(body_children):
        tag = _elem_tag(child)
        if tag == 'p' and not _elem_text(child) and not _has_page_break(child) and not _has_image(child) and not _has_section_break(child):
            if run_start is None:
                run_start = i
            continue
        if run_start is not None:
            count = i - run_start
            if count >= 2:
                runs_to_collapse.append((run_start, i))
            run_start = None

    if run_start is not None:
        count = len(body_children) - run_start
        if count >= 2:
            runs_to_collapse.append((run_start, len(body_children)))

    # Remove excess empty paragraphs (keep 1 at the start of each run)
    body = doc.element.body
    removed = 0
    for start, end in runs_to_collapse:
        keep_first = body_children[start]
        for k in range(start + 1, end):
            if body_children[k] in body:
                body.remove(body_children[k])
                removed += 1

    print(f'Collapsed empty paragraphs: removed {removed}, '
          f'affected {len(runs_to_collapse)} region(s).')


def _add_toc_page_break(doc):
    """Add a page break between the TOC and the first certificate,
    ensuring the TOC stays on its own page.
    Skips if a section break already separates the TOC from certificates."""
    body = list(doc.element.body)

    for i, child in enumerate(body):
        if _elem_tag(child) != 'p':
            continue
        if _elem_text(child) != '产品合格证':
            continue

        # Check if a section break already exists nearby (within 5 paras
        # before the cert title). If so, the section break already provides
        # page separation — no additional page break needed.
        has_nearby_section = False
        for k in range(max(0, i - 5), i):
            if _elem_tag(body[k]) == 'p':
                pPr = body[k].find(qn('w:pPr'))
                if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                    has_nearby_section = True
                    break

        if has_nearby_section:
            print('Section break already separates TOC from certificates — skipping extra page break.')
            return

        # No nearby section break — add a page break before the first cert title
        for j in range(i - 1, -1, -1):
            if _elem_tag(body[j]) == 'p':
                prev = body[j]
                runs = prev.findall(qn('w:r'))
                if runs:
                    last_run = runs[-1]
                else:
                    last_run = OxmlElement('w:r')
                    prev.append(last_run)
                br = OxmlElement('w:br')
                br.set(qn('w:type'), 'page')
                last_run.append(br)
                print('Added page break before first certificate (TOC separation).')
                return
        break


def _compact_visual_tables(doc):
    """Reduce row heights in visual inspection tables for all rows
    except 3, 4, 5, 8, so each report fits on one page."""
    KEEP_ROWS = {3, 4, 5, 8}
    SMALL_HEIGHT = '280'

    body = list(doc.element.body)

    for i, child in enumerate(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'p':
            continue
        text = ''
        for t in child.iter(qn('w:t')):
            if t.text:
                text += t.text
        if '隔震橡胶支座外观质量及尺寸检测报告' not in text:
            continue

        tbl_elem = None
        for j in range(i + 1, min(len(body), i + 5)):
            nc = body[j]
            if (nc.tag.split('}')[-1] if '}' in nc.tag else nc.tag) == 'tbl':
                tbl_elem = nc
                break
        if tbl_elem is None:
            continue

        for ri, tr in enumerate(tbl_elem.findall(qn('w:tr'))):
            if ri in KEEP_ROWS:
                continue

            trPr = tr.find(qn('w:trPr'))
            if trPr is None:
                trPr = OxmlElement('w:trPr')
                tr.insert(0, trPr)
            trHeight = trPr.find(qn('w:trHeight'))
            if trHeight is None:
                trHeight = OxmlElement('w:trHeight')
                trPr.append(trHeight)
            trHeight.set(qn('w:val'), SMALL_HEIGHT)
            trHeight.set(qn('w:hRule'), 'atLeast')

            # Also reduce paragraph spacing in these rows
            for tc in tr.findall(qn('w:tc')):
                for p in tc.findall(qn('w:p')):
                    pPr = p.find(qn('w:pPr'))
                    if pPr is None:
                        pPr = OxmlElement('w:pPr')
                        p.insert(0, pPr)
                    spacing = pPr.find(qn('w:spacing'))
                    if spacing is None:
                        spacing = OxmlElement('w:spacing')
                        pPr.append(spacing)
                    spacing.set(qn('w:before'), '0')
                    spacing.set(qn('w:after'), '0')
                    spacing.set(qn('w:line'), '240')
                    spacing.set(qn('w:lineRule'), 'auto')

    print('Compacted visual inspection table rows.')


def _build_fill_map_for_models(data, product_dir, model_indices, product_type='isolation_bearing'):
    """Build fill_map with data sliced to specific model indices."""
    products = data.get('product_list', [])
    mech = data.get('mechanical_data', [])
    vis = data.get('visual_data', [])

    slim = {
        'project_info': data.get('project_info', {}),
        'product_list': [products[i] for i in model_indices if i < len(products)],
        'mechanical_data': [mech[i] for i in model_indices if i < len(mech)],
        'visual_data': [vis[i] for i in model_indices if i < len(vis)],
    }
    return build_fill_map(slim, product_dir, product_type)


def _needs_hysteresis_curve(excel_data):
    """Return whether the viscous damper report should include hysteresis curve pages."""
    project_info = excel_data.get('project_info', {})
    field_name = '\u662f\u5426\u9700\u8981\u6ede\u56de\u66f2\u7ebf'
    value = str(project_info.get(field_name, '')).strip().lower()
    if not value:
        return True

    false_values = {
        '\u5426',
        '\u4e0d\u9700\u8981',
        '\u4e0d\u8981',
        '\u65e0',
        '0',
        'false',
        'no',
        'n',
    }
    return value not in false_values


def _parse_numeric(value, default):
    if value is None:
        return default
    match = re.search(r'-?\d+(?:\.\d+)?', str(value))
    if not match:
        return default
    try:
        return float(match.group(0))
    except ValueError:
        return default


def _build_hysteresis_curve_images(excel_data, output_dir):
    from core.hysteresis_curve import (
        CURVE_TEMPLATE_KEY,
        DESIGN_DISPLACEMENT_KEY,
        MAX_FORCE_KEY,
        MODEL_KEY,
        choose_curve_template_name,
        generate_hysteresis_curve,
    )
    import tempfile

    os.makedirs(output_dir, exist_ok=True)
    temp_dir = tempfile.mkdtemp(prefix='hysteresis_curves_', dir=output_dir)
    mechanical = excel_data.get('mechanical_data', [])
    products = excel_data.get('product_list', [])
    mech_by_model = {str(row.get(MODEL_KEY, '')): row for row in mechanical}

    image_paths = []
    used_template_names = set()
    for idx, product in enumerate(products):
        model = str(product.get(MODEL_KEY, ''))
        mech_row = dict(mech_by_model.get(model, {}))
        mech_row.setdefault(MODEL_KEY, model)
        mech_row.setdefault(MAX_FORCE_KEY, 300)
        mech_row.setdefault(DESIGN_DISPLACEMENT_KEY, 30)
        max_force = _parse_numeric(mech_row.get(MAX_FORCE_KEY), 300)
        displacement = _parse_numeric(mech_row.get(DESIGN_DISPLACEMENT_KEY), 30)
        template_name = choose_curve_template_name(
            max_force, displacement, used_template_names)
        if template_name:
            mech_row[CURVE_TEMPLATE_KEY] = template_name
            used_template_names.add(template_name)
        image_path = os.path.join(temp_dir, f'hysteresis_{idx + 1}.png')
        generate_hysteresis_curve(mech_row, image_path)
        image_paths.append(image_path)

    return temp_dir, image_paths


def generate_report(template_path, excel_data, output_dir,
                    output_name='报告', product_dir=None,
                    material_manager=None, selected_cert_ids=None,
                    product_type='isolation_bearing'):
    """Generate a report by composing section templates.

    Each section is a standalone .docx file under templates/.
    Code only replaces {{FIELD_XXX}} placeholders — zero format modification.
    Sections are stitched together with docxcompose.
    """
    import json
    from core.doc_composer import compose_report

    if product_dir is None:
        product_dir = os.path.join(
            os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
            'products', 'isolation_bearing'
        )

    config_path = os.path.join(product_dir, 'config.json')
    with open(config_path, 'r', encoding='utf-8') as f:
        config = json.load(f)

    if product_type == 'isolation_bearing_embedded_parts':
        excel_data = _prepare_isolation_embedded_data(
            excel_data, product_dir, material_manager)

    sections = config.get('sections', [])
    if not sections:
        raise ValueError(f'No sections defined in {config_path}')

    templates_dir = os.path.join(product_dir, 'templates')
    products = excel_data.get('product_list', [])
    model_count = len(products)
    mech_per_page = config.get('max_mech_models_per_page', 4)
    include_hysteresis_curve = (
        product_type == 'viscous_damper' and _needs_hysteresis_curve(excel_data)
    )
    curve_temp_dir = None
    curve_image_paths = []
    if include_hysteresis_curve:
        curve_temp_dir, curve_image_paths = _build_hysteresis_curve_images(
            excel_data, output_dir)

    section_specs = []
    for sec in sections:
        section_file = sec['file']
        if (
            product_type == 'viscous_damper'
            and section_file == '08_mechanical.docx'
            and not _needs_hysteresis_curve(excel_data)
        ):
            section_file = '08-1_mechanical.docx'

        tmpl = os.path.join(templates_dir, section_file)
        repeat_spec = sec.get('repeat', 1)
        per_model = sec.get('per_model', False)
        batch_size = sec.get('batch_size', 0)

        if product_type == 'isolation_bearing_embedded_parts':
            if repeat_spec == 'certificate_pages':
                rows = excel_data.get('embedded_certificate_rows', [])
                batches = [rows[i:i + 13] for i in range(0, len(rows), 13)] or [[]]
                fill_maps = [
                    _build_isolation_embedded_cert_fill_map(excel_data, batch)
                    for batch in batches
                ]
                for fill_map in fill_maps:
                    section_specs.append({
                        'template': tmpl,
                        'repeat': 1,
                        'fill_map': fill_map,
                        'page_break_before': sec.get('page_break_before', True),
                    })
                continue
            if repeat_spec == 'anchor_count':
                items = excel_data.get('embedded_anchor_components', [])
                for item in items:
                    section_specs.append({
                        'template': tmpl,
                        'repeat': 1,
                        'fill_map': _build_isolation_embedded_anchor_fill_map(
                            excel_data, item),
                        'page_break_before': sec.get('page_break_before', True),
                    })
                continue
            if repeat_spec == 'plate_count':
                items = excel_data.get('embedded_plate_components', [])
                for item in items:
                    section_specs.append({
                        'template': tmpl,
                        'repeat': 1,
                        'fill_map': _build_isolation_embedded_plate_fill_map(
                            excel_data, item),
                        'page_break_before': sec.get('page_break_before', True),
                    })
                continue

        if product_type == 'friction_pendulum' and section_file == '06_mech_detail.docx':
            serial_rows = _expand_friction_serials(products)
            detail_batches = [
                serial_rows[i:i + 16] for i in range(0, len(serial_rows), 16)
            ] or [[]]
            section_specs.append({
                'template': tmpl,
                'repeat': len(detail_batches),
                'fill_maps': [_build_friction_detail_fill_map(batch) for batch in detail_batches],
                'page_break_after': sec.get('page_break_after', False),
                'page_break_before': sec.get('page_break_before', True),
            })
            continue

        # Resolve repeat count
        if repeat_spec == 'model_count':
            repeat = max(1, model_count)
        elif repeat_spec == 'mech_pages':
            repeat = max(1, (model_count + mech_per_page - 1) // mech_per_page)
        else:
            repeat = int(repeat_spec)

        if repeat > 1 and per_model:
            fill_maps = [
                _build_fill_map_for_models(excel_data, product_dir, [i], product_type)
                for i in range(model_count)
            ]
            section_specs.append({
                'template': tmpl, 'repeat': repeat, 'fill_maps': fill_maps,
                'page_break_after': sec.get('page_break_after', False),
                'page_break_before': sec.get('page_break_before', True),
                'image_paths': curve_image_paths if section_file == '08_mechanical.docx' else None,
            })
        elif repeat > 1 and batch_size:
            fill_maps = []
            batch_counts = []
            for i in range(repeat):
                indices = list(range(i * batch_size, min((i + 1) * batch_size, model_count)))
                if indices:
                    fill_maps.append(
                        _build_fill_map_for_models(excel_data, product_dir, indices, product_type))
                    batch_counts.append(len(indices))
            section_specs.append({
                'template': tmpl, 'repeat': len(fill_maps), 'fill_maps': fill_maps,
                'page_break_after': sec.get('page_break_after', False),
                'page_break_before': sec.get('page_break_before', True),
                'batch_counts': batch_counts,
                'image_paths': curve_image_paths if section_file == '08_mechanical.docx' else None,
            })
        else:
            fill_map = _build_fill_map_for_models(
                excel_data, product_dir, list(range(model_count)), product_type)
            spec = {
                'template': tmpl, 'repeat': repeat, 'fill_map': fill_map,
                'page_break_after': sec.get('page_break_after', False),
                'page_break_before': sec.get('page_break_before', True),
                'image_paths': curve_image_paths[:1] if section_file == '08_mechanical.docx' else None,
            }
            if batch_size:
                spec['batch_counts'] = [model_count]
            section_specs.append(spec)

    if product_type == 'isolation_bearing_embedded_parts':
        for spec in section_specs[1:]:
            spec['stable_page_break_before'] = True

    # Output path
    project_name = excel_data.get('project_info', {}).get('项目名称', '未命名项目')
    report_date = excel_data.get('project_info', {}).get('制造日期', '未知日期')
    safe_date = _normalize_report_month(report_date)

    output_subdir = os.path.join(output_dir, project_name, safe_date)
    os.makedirs(output_subdir, exist_ok=True)

    base_path = os.path.join(output_subdir, f'{output_name}.docx')
    output_path = base_path
    version = 1
    while os.path.exists(output_path):
        output_path = os.path.join(output_subdir, f'{output_name}_V{version}.docx')
        version += 1

    try:
        if product_type == 'isolation_bearing_embedded_parts':
            auto_ids = excel_data.get('_auto_material_cert_ids', [])
            if auto_ids:
                selected_cert_ids = auto_ids
        compose_report(section_specs, output_path, excel_data, material_manager, selected_cert_ids)
    finally:
        if curve_temp_dir:
            import shutil
            shutil.rmtree(curve_temp_dir, ignore_errors=True)

    return output_path

def _replace_all(doc, fill_map):
    """Replace all {{FIELD_XXX}} placeholders in a document."""

    def replace_in_paragraphs(paragraphs):
        for para in paragraphs:
            for run in para.runs:
                for placeholder, value in fill_map.items():
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)

    replace_in_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_paragraphs(cell.paragraphs)
    for section in doc.sections:
        replace_in_paragraphs(section.header.paragraphs)
        replace_in_paragraphs(section.footer.paragraphs)
