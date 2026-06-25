import os
import re
import copy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxcompose.composer import Composer


def _count_fields(doc):
    text = ''
    for p in doc.paragraphs:
        text += ' '.join(r.text or '' for r in p.runs) + '\n'
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                text += c.text + '\n'
    return len(re.findall(r'\{\{[^}]+\}\}', text))


def _replace_text_nodes(root, fill_map):
    """Replace placeholders even when Word splits them across runs."""
    if not fill_map:
        return

    for ph, val in fill_map.items():
        replacement = str(val)
        while True:
            text_nodes = [node for node in root.iter(qn('w:t')) if node.text is not None]
            full_text = ''.join(node.text or '' for node in text_nodes)
            start = full_text.find(ph)
            if start < 0:
                break
            end = start + len(ph)

            offsets = [0]
            for node in text_nodes:
                offsets.append(offsets[-1] + len(node.text or ''))

            placed = False
            for idx, node in enumerate(text_nodes):
                node_start = offsets[idx]
                node_end = offsets[idx + 1]
                if node_end <= start or node_start >= end:
                    continue

                text = node.text or ''
                prefix = text[:max(0, start - node_start)]
                suffix = text[max(0, end - node_start):]
                if not placed:
                    node.text = prefix + replacement + suffix
                    placed = True
                else:
                    node.text = prefix + suffix


def _replace_all(doc, fill_map):
    _replace_text_nodes(doc.element.body, fill_map)


def _clone_body_content(doc, times):
    if times <= 0:
        return
    body = doc.element.body
    original_children = list(body)
    last_child = original_children[-1]
    for _ in range(times):
        for child in original_children:
            cloned = copy.deepcopy(child)
            last_child.addnext(cloned)
            last_child = cloned


def _trim_mech_table_rows(doc, keep_models):
    if keep_models >= 4:
        return
    for table in doc.tables:
        trs = table._tbl.findall(qn('w:tr'))
        if len(trs) != 16:
            continue
        target = keep_models * 4
        for i in range(len(trs) - 1, target - 1, -1):
            table._tbl.remove(trs[i])


def _replace_inline_pictures(doc, image_paths):
    if not image_paths:
        return

    blips = list(doc.element.body.iter(qn('a:blip')))
    for blip, image_path in zip(blips, image_paths):
        if not image_path:
            continue
        r_id, _ = doc.part.get_or_add_image(image_path)
        blip.set(qn('r:embed'), r_id)


def process_section(template_path, fill_map, repeat=1, fill_maps=None,
                   between_pages=False, batch_counts=None, image_paths=None,
                   page_break_every=0):
    doc = Document(template_path)

    if repeat > 1:
        _clone_body_content(doc, repeat - 1)

        if fill_maps and len(fill_maps) == repeat:
            all_children = list(doc.element.body)
            group_size = len(all_children) // repeat
            for i in range(repeat):
                start = i * group_size
                end = start + group_size
                fm = fill_maps[i]
                for j in range(start, min(end, len(all_children))):
                    child = all_children[j]
                    _replace_text_nodes(child, fm)
        else:
            _replace_all(doc, fill_map)

        if (between_pages or page_break_every) and repeat > 1:
            all_children = list(doc.element.body)
            group_size = len(all_children) // repeat
            for g in range(1, repeat):
                if page_break_every and g % page_break_every != 0:
                    continue
                insert_before = all_children[g * group_size]
                pb_para = _make_page_break_para()
                insert_before.addprevious(pb_para)

        if batch_counts:
            _apply_batch_trimming(doc, batch_counts)
    else:
        _replace_all(doc, fill_map)
        if batch_counts and batch_counts[0] < 4:
            _trim_mech_table_rows(doc, batch_counts[0])

    _replace_inline_pictures(doc, image_paths)

    return doc


def _make_page_break_para():
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    return p


def _set_page_break_before_first_content(doc):
    """Start a section on a new page without an overflow-prone break paragraph."""
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            p_pr = child.find(qn('w:pPr'))
            if p_pr is None:
                p_pr = OxmlElement('w:pPr')
                child.insert(0, p_pr)
            if p_pr.find(qn('w:pageBreakBefore')) is None:
                p_pr.append(OxmlElement('w:pageBreakBefore'))
            return
        if tag == 'tbl':
            paragraph = OxmlElement('w:p')
            p_pr = OxmlElement('w:pPr')
            p_pr.append(OxmlElement('w:pageBreakBefore'))
            paragraph.append(p_pr)
            child.addprevious(paragraph)
            return


def _absorb_empty_pb_paras(doc):
    """Post-compose step: find empty paragraphs that only contain a page break,
    migrate the page break into the preceding content paragraph's last run,
    then remove the empty paragraph. This prevents blank pages without moving
    breaks across tables or image-only pages."""
    body = doc.element.body
    children = list(body)
    modified = False

    for child in list(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'p':
            continue
        text = ''.join(t.text or '' for t in child.iter(qn('w:t')))
        if text.strip():
            continue
        has_pb = any(br.get(qn('w:type')) == 'page' for br in child.iter(qn('w:br')))
        if not has_pb:
            continue
        # Has image? keep it
        if any(child.findall('.//' + qn(t_)) for t_ in ('w:drawing', 'w:pict', 'w:object')):
            continue
        # Has section break? keep it
        pPr = child.find(qn('w:pPr'))
        if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
            continue

        child_pos = children.index(child)
        if child_pos == 0:
            continue

        # Only absorb into the immediately preceding paragraph. Searching
        # farther back can move a chapter break ahead of image-only pages
        # such as business licenses or material certificates.
        previous_para = children[child_pos - 1]
        ptag = previous_para.tag.split('}')[-1] if '}' in previous_para.tag else previous_para.tag
        if ptag != 'p':
            continue
        prev_has_text = ''.join(t.text or '' for t in previous_para.iter(qn('w:t'))).strip()
        prev_has_image = any(previous_para.findall('.//' + qn(t_))
                             for t_ in ('w:drawing', 'w:pict', 'w:object'))
        if not prev_has_text and not prev_has_image:
            continue

        # Get the last run of the previous paragraph
        runs = previous_para.findall(qn('w:r'))
        if runs:
            last_run = runs[-1]
        else:
            last_run = OxmlElement('w:r')
            previous_para.append(last_run)
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        last_run.append(br)
        body.remove(child)
        modified = True

    if modified:
        print('Absorbed empty page-break paragraphs into preceding content.')


def _apply_batch_trimming(doc, batch_counts):
    batch_idx = 0
    for table in doc.tables:
        trs = table._tbl.findall(qn('w:tr'))
        if len(trs) != 16:
            continue
        if batch_idx >= len(batch_counts):
            break
        count = batch_counts[batch_idx]
        if count < 4:
            target = count * 4
            for ri in range(len(trs) - 1, target - 1, -1):
                table._tbl.remove(trs[ri])
        batch_idx += 1


def _make_section_break_para(orientation):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    sectPr = OxmlElement('w:sectPr')
    pgSz = OxmlElement('w:pgSz')
    if orientation == 'landscape':
        pgSz.set(qn('w:w'), '16839')
        pgSz.set(qn('w:h'), '11906')
        pgSz.set(qn('w:orient'), 'landscape')
    else:
        pgSz.set(qn('w:w'), '11906')
        pgSz.set(qn('w:h'), '16839')
    sectPr.append(pgSz)
    pPr.append(sectPr)
    p.append(pPr)
    return p


def _template_orientation(template_path):
    for s in Document(template_path).sections:
        if s.page_width > s.page_height:
            return 'landscape'
    return 'portrait'


def compose_report(section_specs, output_path, excel_data,
                   material_mgr=None, selected_cert_ids=None):
    specs = [s for s in section_specs if s.get('repeat', 1) > 0]
    if not specs:
        return None

    first = specs[0]
    master = process_section(
        first['template'], first.get('fill_map', {}),
        first.get('repeat', 1), first.get('fill_maps'),
        first.get('page_break_after', False), first.get('batch_counts'),
        first.get('image_paths'), first.get('page_break_every', 0))

    composer = Composer(master)

    previous_spec = first
    for spec in specs[1:]:
        section_doc = process_section(
            spec['template'], spec.get('fill_map', {}),
            spec.get('repeat', 1), spec.get('fill_maps'),
            spec.get('page_break_after', False), spec.get('batch_counts'),
            spec.get('image_paths'), spec.get('page_break_every', 0))

        if previous_spec.get('page_break_after', False) or spec.get('page_break_before', True):
            if spec.get('stable_page_break_before', False):
                _set_page_break_before_first_content(section_doc)
            else:
                master.add_page_break()
        composer.append(section_doc)
        previous_spec = spec

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    composer.save(output_path)

    result = Document(output_path)

    # Migrate page breaks from empty paragraphs into preceding content
    _absorb_empty_pb_paras(result)

    _fix_orientations(result, specs)
    _direct_fill_embedded_cert(result, excel_data)
    _direct_fill_mech_detail(result, excel_data)
    _direct_fill_visual(result, excel_data)
    _merge_friction_mech_model_cells(result)

    if material_mgr and selected_cert_ids:
        matched = material_mgr.get_certs_by_ids(selected_cert_ids)
        if matched:
            material_mgr.insert_material_certs(result, matched)

    _remove_final_paragraph_section_break(result)

    result.save(output_path)

    unfilled = _count_fields(result)
    if unfilled > 0:
        print(f'WARNING: {unfilled} unfilled placeholders in output!')

    return output_path


def _get_anchor_text(template_path):
    doc = Document(template_path)
    candidates = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t or len(t) < 3:
            continue
        if '{{' in t and '}}' in t:
            continue
        candidates.append(t)
    title_markers = ('报告', '合格证', '目录', '试验曲线', '质量证明书')
    boilerplate_prefixes = (
        '四川融海运通抗震科技有限责任公司',
        '地址：',
        '电话：',
    )
    for t in candidates:
        if any(t.startswith(prefix) for prefix in boilerplate_prefixes):
            continue
        if any(marker in t for marker in title_markers):
            return t
    if candidates:
        return candidates[0]
    if doc.tables:
        return doc.tables[0].cell(0, 0).text.strip()
    return ''


def _remove_final_paragraph_section_break(doc):
    """Avoid a trailing blank page caused by a final paragraph sectPr."""
    body = doc.element.body
    for child in reversed(list(body)):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'sectPr':
            continue
        sect_pr = child.find('.//w:pPr/w:sectPr', child.nsmap)
        if sect_pr is not None:
            body_sect = body.find(qn('w:sectPr'))
            if body_sect is not None:
                body.remove(body_sect)
            body.append(copy.deepcopy(sect_pr))
            sect_pr.getparent().remove(sect_pr)
        break


def _fix_orientations(doc, specs):
    anchors = [_get_anchor_text(s['template']) for s in specs]
    orientations = [_template_orientation(s['template']) for s in specs]

    if len(set(orientations)) <= 1:
        return

    body = doc.element.body
    children = list(body)

    spec_positions = {}
    search_from = 0
    for spec_idx, (anchor, orient) in enumerate(zip(anchors, orientations)):
        if not anchor:
            continue
        found = False
        for i in range(search_from, len(children)):
            child = children[i]
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag != 'p':
                continue
            text = ''.join(t.text or '' for t in child.iter(qn('w:t'))).strip()
            if text == anchor:
                spec_positions[spec_idx] = i
                search_from = i + 1
                found = True
                break
        if not found:
            return

    ordered = sorted(spec_positions.items())

    for i in range(len(ordered) - 1):
        si, body_idx = ordered[i]
        next_si, next_body_idx = ordered[i + 1]
        if orientations[si] != orientations[next_si]:
            insert_idx = next_body_idx
            for k in range(next_body_idx - 1, body_idx, -1):
                child = children[k]
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if tag != 'p':
                    break
                has_pb = any(br.get(qn('w:type')) == 'page'
                             for br in child.iter(qn('w:br')))
                if has_pb:
                    insert_idx = k + 1
                    break
            ref = children[insert_idx]
            ref.addprevious(_make_section_break_para(orientations[si]))

    last_si = ordered[-1][0]
    last_break = _make_section_break_para(orientations[last_si])
    body_sect = doc.element.body.find(qn('w:sectPr'))
    if body_sect is not None:
        body_sect.addprevious(last_break)
    else:
        doc.element.body.append(last_break)


def _set_cell_text(cell_elem, text):
    all_t = list(cell_elem.iter(qn('w:t')))
    for t in all_t:
        t.text = ''
    if all_t:
        all_t[0].text = text


def _merge_friction_mech_model_cells(doc):
    """Merge repeated model cells in friction pendulum mechanical detail pages."""
    for table in doc.tables:
        if len(table.rows) != 19 or len(table.columns) != 7:
            continue

        trs = table._tbl.findall(qn('w:tr'))
        if len(trs) != 19:
            continue
        header_cells = trs[0].findall(qn('w:tc'))
        if len(header_cells) != 3:
            continue
        second_row_cells = trs[1].findall(qn('w:tc'))
        if len(second_row_cells) != 7:
            continue

        start_idx = None
        current_model = ''
        for row_idx in range(3, len(trs) + 1):
            model = ''
            if row_idx < len(trs):
                cells = trs[row_idx].findall(qn('w:tc'))
                if cells:
                    model = ''.join(t.text or '' for t in cells[0].iter(qn('w:t'))).strip()

            if row_idx < len(trs) and model and model == current_model:
                continue

            if start_idx is not None and row_idx - start_idx > 1:
                _apply_vertical_merge(trs, start_idx, row_idx - 1, 0, current_model)

            if row_idx < len(trs) and model:
                start_idx = row_idx
                current_model = model
            else:
                start_idx = None
                current_model = ''


def _apply_vertical_merge(rows, start_idx, end_idx, col_idx, text):
    for row_idx in range(start_idx, end_idx + 1):
        cells = rows[row_idx].findall(qn('w:tc'))
        if col_idx >= len(cells):
            continue
        cell = cells[col_idx]
        tc_pr = cell.find(qn('w:tcPr'))
        if tc_pr is None:
            tc_pr = OxmlElement('w:tcPr')
            cell.insert(0, tc_pr)
        v_merge = tc_pr.find(qn('w:vMerge'))
        if v_merge is None:
            v_merge = OxmlElement('w:vMerge')
            tc_pr.append(v_merge)
        if row_idx == start_idx:
            v_merge.set(qn('w:val'), 'restart')
            _set_cell_text(cell, text)
            if start_idx > 3:
                _set_cell_border(cell, 'top')
        else:
            v_merge.set(qn('w:val'), 'continue')
            _set_cell_text(cell, '')
            if row_idx == end_idx:
                _set_cell_border(cell, 'bottom')


def _set_cell_border(cell, edge):
    tc_pr = cell.find(qn('w:tcPr'))
    if tc_pr is None:
        tc_pr = OxmlElement('w:tcPr')
        cell.insert(0, tc_pr)
    borders = tc_pr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    border = borders.find(qn(f'w:{edge}'))
    if border is None:
        border = OxmlElement(f'w:{edge}')
        borders.append(border)
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), '4')
    border.set(qn('w:space'), '0')
    border.set(qn('w:color'), '000000')


def _compute_judgment(val):
    if val is None:
        return '/'
    s = str(val).strip().replace('%', '')
    if s in ('/', '0', '', 'None', '.'):
        return '/'
    try:
        float(s)
        return '合格'
    except ValueError:
        return '/'


def _direct_fill_embedded_cert(doc, excel_data):
    """Fill embedded-damper-parts certificate rows from the product list."""
    products = excel_data.get('product_list', [])
    if not products:
        return

    for table in doc.tables:
        if len(table.rows) < 2 or len(table.columns) < 4:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells[:4]]
        if headers != ['产品名称', '规格型号（mm）', '生产日期', '数量（件）']:
            continue

        tbl = table._tbl
        rows = list(tbl.findall(qn('w:tr')))
        data_rows = rows[1:]

        while len(data_rows) < len(products):
            cloned = copy.deepcopy(data_rows[-1])
            tbl.append(cloned)
            data_rows.append(cloned)

        for extra in data_rows[len(products):]:
            tbl.remove(extra)

        rows = list(tbl.findall(qn('w:tr')))[1:1 + len(products)]
        for row, product in zip(rows, products):
            cells = row.findall(qn('w:tc'))
            values = [
                product.get('产品名称', '预埋件'),
                product.get('产品型号', ''),
                product.get('生产日期', ''),
                product.get('数量', ''),
            ]
            for cell, value in zip(cells[:4], values):
                _set_cell_text(cell, str(value))
        return


def _direct_fill_visual(doc, excel_data):
    products = excel_data.get('product_list', [])
    project = excel_data.get('project_info', {})
    body = doc.element.body
    body_children = list(body)

    vis_page = 0
    for i, child in enumerate(body_children):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'p':
            continue
        text = ''.join(t.text or '' for t in child.iter(qn('w:t')))
        if '隔震橡胶支座外观质量及尺寸检测报告' not in text:
            continue

        tbl_elem = None
        for j in range(i + 1, min(len(body_children), i + 5)):
            nc = body_children[j]
            if (nc.tag.split('}')[-1] if '}' in nc.tag else nc.tag) == 'tbl':
                tbl_elem = nc
                break
        if tbl_elem is None:
            continue

        trs = tbl_elem.findall(qn('w:tr'))
        if len(trs) < 21:
            continue
        if vis_page >= len(products):
            vis_page += 1
            continue

        product = products[vis_page]
        model_name = product.get('产品型号', '')

        cells_r0 = trs[0].findall(qn('w:tc'))
        if len(cells_r0) > 2:
            _set_cell_text(cells_r0[2], str(project.get('项目名称', '')))

        cells_r1 = trs[1].findall(qn('w:tc'))
        if len(cells_r1) > 2:
            _set_cell_text(cells_r1[2], str(product.get('数量', '')))

        cells_r3 = trs[3].findall(qn('w:tc'))
        old_model = ''
        if len(cells_r3) > 3:
            for t in cells_r3[3].iter(qn('w:t')):
                if t.text:
                    old_model += t.text
        old_model = old_model.strip()
        if len(cells_r3) > 3:
            _set_cell_text(cells_r3[3], model_name)

        if old_model:
            last_tr = trs[-1]
            for cell in last_tr.findall(qn('w:tc')):
                all_ts = list(cell.iter(qn('w:t')))
                combined = ''.join((t.text or '') for t in all_ts)
                pos = combined.find(old_model)
                if pos < 0:
                    continue
                end_pos = pos + len(old_model)
                done = False
                offset = 0
                for t in all_ts:
                    txt = t.text or ''
                    rlen = len(txt)
                    if rlen == 0:
                        offset += 0
                        continue
                    r_start = offset
                    r_end = offset + rlen
                    offset += rlen
                    if r_end <= pos or r_start >= end_pos:
                        continue
                    prefix = txt[:max(0, pos - r_start)]
                    suffix = txt[max(0, end_pos - r_start):]
                    if not done:
                        t.text = prefix + model_name + suffix
                        done = True
                    else:
                        t.text = prefix + suffix
                break

        vis_page += 1


def _direct_fill_mech_detail(doc, excel_data):
    products = excel_data.get('product_list', [])
    mechanical = excel_data.get('mechanical_data', [])
    mech_by_model = {m['产品型号']: m for m in mechanical}

    body = doc.element.body
    body_children = list(body)
    model_idx = 0

    JUDGMENT_KEYS = [
        '设计压缩刚度(KN/mm)', '水平等效刚度(KN/mm)',
        '屈服后刚度(KN/mm)', '屈服力(KN)', '等效阻尼比(%)'
    ]

    for i, child in enumerate(body_children):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'p':
            continue
        text = ''.join(t.text or '' for t in child.iter(qn('w:t')))
        if '隔震支座力学性能检验报告' not in text:
            continue

        tbl_elem = None
        for j in range(i + 1, min(len(body_children), i + 5)):
            nc = body_children[j]
            if (nc.tag.split('}')[-1] if '}' in nc.tag else nc.tag) == 'tbl':
                tbl_elem = nc
                break
        if tbl_elem is None:
            continue

        trs = tbl_elem.findall(qn('w:tr'))
        if len(trs) > 16:
            continue
        group_count = len(trs) // 4
        for g in range(group_count):
            data_row_idx = g * 4 + 3
            if data_row_idx >= len(trs) or model_idx >= len(products):
                break
            mech = mech_by_model.get(products[model_idx].get('产品型号', ''), {})
            cells = trs[data_row_idx].findall(qn('w:tc'))
            for col_idx in range(2, min(7, len(cells))):
                key = JUDGMENT_KEYS[col_idx - 2]
                _set_cell_text(cells[col_idx], _compute_judgment(mech.get(key, '/')))
            model_idx += 1
