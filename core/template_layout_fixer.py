"""
Add page breaks to the prepared template at key positions
to ensure proper pagination in the generated report.

Rules:
- Certificates (Tables 0-5): NO page breaks — natural flow
- Mechanical report (Table 6): page break BEFORE the company name
- Visual inspection reports (Tables 9-14): page break BEFORE the company name
"""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _add_page_break_to_last_run(paragraph):
    """Append a page break to the last run of a paragraph.
    If the paragraph has no runs, create an empty run first."""
    if not paragraph.runs:
        paragraph.add_run('')
    last_run = paragraph.runs[-1]
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    last_run._element.append(br)
    return True


def _find_para_element_before_company(doc, body_items, table_item_idx):
    """
    Starting from a table's position, scan backwards to find the
    company-name paragraph, then return the element of the paragraph
    just before it (where the page break should go).
    """
    # Scan backwards from table to find company-name paragraph
    company_pos = None
    for j in range(table_item_idx - 1, -1, -1):
        if body_items[j][0] == 'p':
            text = ''
            for t in body_items[j][2].iter(qn('w:t')):
                if t.text:
                    text += t.text
            text = text.strip()
            # The company name paragraph appears as section headers
            if '四川融海运通' in text:
                company_pos = j
                break

    if company_pos is None:
        return None

    # Find the paragraph just before the company name paragraph
    for j in range(company_pos - 1, -1, -1):
        if body_items[j][0] == 'p':
            return body_items[j][2]
    return None


def _get_para_text(elem):
    """Get text from a paragraph XML element."""
    text = ''
    for t in elem.iter(qn('w:t')):
        if t.text:
            text += t.text
    return text.strip()[:60]


def fix_template_pagination(template_path, output_path=None):
    """
    Add page breaks before report sections (not within certificate pages).
    """
    if output_path is None:
        output_path = template_path

    doc = Document(template_path)
    body = doc.element.body

    # Build ordered list of body children
    body_items = []
    para_idx = 0
    tbl_idx = 0
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            body_items.append(('p', para_idx, child))
            para_idx += 1
        elif tag == 'tbl':
            body_items.append(('tbl', tbl_idx, child))
            tbl_idx += 1

    # Target tables that start new report sections
    target_tables = [6] + list(range(9, 15))
    breaks_added = 0

    for i, (etype, idx, elem) in enumerate(body_items):
        if etype != 'tbl' or idx not in target_tables:
            continue

        prev_para_elem = _find_para_element_before_company(doc, body_items, i)
        if prev_para_elem is None:
            print(f'Table[{idx}]: no company-name paragraph found, skipping')
            continue

        # Find matching paragraph object in doc.paragraphs
        for para in doc.paragraphs:
            if para._element is prev_para_elem:
                prev_text = _get_para_text(prev_para_elem)
                print(f'Table[{idx}]: page break after para "{prev_text}"')
                if _add_page_break_to_last_run(para):
                    breaks_added += 1
                break

    print(f'Added {breaks_added} page breaks.')
    doc.save(output_path)
    return output_path


if __name__ == '__main__':
    import os
    _base = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    fix_template_pagination(
        os.path.join(_base, 'products', 'isolation_bearing', 'template_prepared.docx')
    )
