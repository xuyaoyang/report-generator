import unittest

from docx import Document
from docx.oxml.ns import qn

from core.doc_composer import _clone_body_content


class DocComposerTest(unittest.TestCase):
    def test_clone_body_content_does_not_duplicate_body_sectpr(self):
        doc = Document()
        doc.add_paragraph('第一页')

        _clone_body_content(doc, 2)

        sectpr_count = sum(
            1 for child in doc.element.body
            if child.tag == qn('w:sectPr')
        )
        self.assertEqual(sectpr_count, 1)
        self.assertEqual(len(doc.paragraphs), 3)


if __name__ == '__main__':
    unittest.main()
