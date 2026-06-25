"""Build section templates from a friction-pendulum embedded-parts report."""
import copy
import os
import shutil
import sys
import tempfile
import zipfile
from posixpath import basename

from docx import Document
from docx.oxml.ns import qn
from lxml import etree


PRODUCT_DIR = os.path.dirname(os.path.abspath(__file__))
ROOT_DIR = os.path.dirname(os.path.dirname(PRODUCT_DIR))
TEMPLATES_DIR = os.path.join(PRODUCT_DIR, 'templates')
REL_NS = 'http://schemas.openxmlformats.org/package/2006/relationships'
OFFICE_REL_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
CONTENT_TYPE_NS = 'http://schemas.openxmlformats.org/package/2006/content-types'
NUMBERING_REL = (
    'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering')
NUMBERING_CONTENT_TYPE = (
    'application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml')
MINIMAL_NUMBERING = b'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'''


def _extract(source, output, start, end, section_index=2):
    doc = Document(source)
    body = doc.element.body
    children = list(body)
    section_index = min(section_index, len(doc.sections) - 1)
    final_section = copy.deepcopy(doc.sections[section_index]._sectPr)
    section_type = final_section.find(qn('w:type'))
    if section_type is None:
        section_type = final_section.makeelement(qn('w:type'))
        final_section.insert(0, section_type)
    section_type.set(qn('w:val'), 'continuous')
    for index in range(len(children) - 1, end, -1):
        body.remove(children[index])
    for index in range(start - 1, -1, -1):
        body.remove(children[index])
    for child in list(body):
        nested_section = child.find('.//' + qn('w:sectPr'))
        if nested_section is not None:
            nested_section.getparent().remove(nested_section)
        if child.tag == qn('w:sectPr'):
            body.remove(child)
    body.append(final_section)
    doc.save(output)


def _prune_unused_document_images(path):
    with zipfile.ZipFile(path, 'r') as source:
        files = {name: source.read(name) for name in source.namelist()}

    document = etree.fromstring(files['word/document.xml'])
    used_ids = set(document.xpath(
        '//@r:embed | //@r:id', namespaces={'r': OFFICE_REL_NS}))
    rel_name = 'word/_rels/document.xml.rels'
    rels = etree.fromstring(files[rel_name])
    for rel in list(rels):
        if rel.get('Type', '').endswith('/image') and rel.get('Id') not in used_ids:
            rels.remove(rel)
    files[rel_name] = etree.tostring(
        rels, xml_declaration=True, encoding='UTF-8', standalone=True)

    referenced_media = set()
    for name, content in files.items():
        if not name.endswith('.rels'):
            continue
        root = etree.fromstring(content)
        for rel in root.findall(f'{{{REL_NS}}}Relationship'):
            target = rel.get('Target', '')
            if 'media/' in target:
                referenced_media.add(basename(target))
    for name in list(files):
        if name.startswith('word/media/') and basename(name) not in referenced_media:
            del files[name]

    fd, temp_path = tempfile.mkstemp(suffix='.docx', dir=os.path.dirname(path))
    os.close(fd)
    try:
        with zipfile.ZipFile(temp_path, 'w', zipfile.ZIP_DEFLATED) as target:
            for name, content in files.items():
                target.writestr(name, content)
        os.replace(temp_path, path)
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def _ensure_numbering_part(path):
    with zipfile.ZipFile(path, 'r') as source:
        files = {name: source.read(name) for name in source.namelist()}
    if 'word/numbering.xml' in files:
        return

    rel_name = 'word/_rels/document.xml.rels'
    rels = etree.fromstring(files[rel_name])
    rel_ids = {rel.get('Id') for rel in rels}
    index = 1
    while f'rId{index}' in rel_ids:
        index += 1
    rel = etree.Element(f'{{{REL_NS}}}Relationship')
    rel.set('Id', f'rId{index}')
    rel.set('Type', NUMBERING_REL)
    rel.set('Target', 'numbering.xml')
    rels.append(rel)
    files[rel_name] = etree.tostring(
        rels, xml_declaration=True, encoding='UTF-8', standalone=True)

    types = etree.fromstring(files['[Content_Types].xml'])
    exists = any(
        item.get('PartName') == '/word/numbering.xml'
        for item in types.findall(f'{{{CONTENT_TYPE_NS}}}Override'))
    if not exists:
        override = etree.Element(f'{{{CONTENT_TYPE_NS}}}Override')
        override.set('PartName', '/word/numbering.xml')
        override.set('ContentType', NUMBERING_CONTENT_TYPE)
        types.append(override)
        files['[Content_Types].xml'] = etree.tostring(
            types, xml_declaration=True, encoding='UTF-8', standalone=True)
    files['word/numbering.xml'] = MINIMAL_NUMBERING

    fd, temp_path = tempfile.mkstemp(suffix='.docx', dir=os.path.dirname(path))
    os.close(fd)
    try:
        with zipfile.ZipFile(temp_path, 'w', zipfile.ZIP_DEFLATED) as target:
            for name, content in files.items():
                target.writestr(name, content)
        os.replace(temp_path, path)
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def _set_paragraph(paragraph, value):
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ''
    else:
        paragraph.add_run(value)


def _set_cell(cell, value):
    paragraph = cell.paragraphs[0]
    _set_paragraph(paragraph, value)
    for extra in cell.paragraphs[1:]:
        _set_paragraph(extra, '')


def _replace_prefixed_paragraph(doc, prefix, value):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            _set_paragraph(paragraph, value)
            return


def _replace_contains_paragraph(doc, text, value):
    for paragraph in doc.paragraphs:
        if text in paragraph.text:
            _set_paragraph(paragraph, value)
            return


def _prepare_cover(path):
    doc = Document(path)
    _replace_prefixed_paragraph(doc, '四川融海运通', '{{IBE_SUPPLIER}}')
    _replace_contains_paragraph(doc, '预埋组件出厂资料', '{{IBE_TITLE}}')
    _replace_prefixed_paragraph(doc, '工程名称', '工程名称：{{IBE_PROJECT}}')
    _replace_prefixed_paragraph(doc, '产品供应商', '产品供应商：{{IBE_SUPPLIER}}')
    _replace_prefixed_paragraph(doc, '生产地址', '生产地址：{{IBE_ADDRESS}}')
    _replace_prefixed_paragraph(doc, '电', '电    话：{{IBE_PHONE}}')
    _replace_prefixed_paragraph(doc, '报告日期', '报告日期：{{IBE_DATE}}')
    doc.save(path)


def _prepare_cert(path):
    doc = Document(path)
    table = doc.tables[0]
    for index in range(1, min(14, len(table.rows))):
        for col, field in enumerate(('NAME', 'SPEC', 'DATE', 'QTY')):
            _set_cell(table.cell(index, col), f'{{{{IBE_CERT_{index}_{field}}}}}')
    doc.save(path)


def _prepare_report_common(doc, title_placeholder):
    _replace_prefixed_paragraph(doc, '四川融海运通', '{{IBE_SUPPLIER}}')
    for paragraph in doc.paragraphs:
        if '出厂检测报告' in paragraph.text:
            _set_paragraph(paragraph, title_placeholder)
            break
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith('主检'):
            _set_paragraph(
                paragraph,
                '主检：{{IBE_INSPECTOR}}                 审核：{{IBE_REVIEWER}}                  签发：{{IBE_ISSUER}}')
            break


def _prepare_plate(path):
    doc = Document(path)
    _prepare_report_common(doc, '{{IBE_COMPONENT_NAME}}出厂检测报告')
    table = doc.tables[0]
    for cell in table.rows[0].cells:
        _set_cell(cell, '规格：{{IBE_PLATE_SPEC}}')
    for cell in table.rows[2].cells:
        text = cell.text
        if '±2' in text:
            _set_cell(cell, '{{IBE_PLATE_SIDE_STANDARD}}')
        elif '±0.5' in text:
            _set_cell(cell, '{{IBE_PLATE_THICKNESS_STANDARD}}')
        elif '±0.8' in text:
            _set_cell(cell, '{{IBE_PLATE_HOLE_STANDARD}}')
    for row in range(6, 9):
        _set_cell(table.cell(row, 0), '热轧钢板（{{IBE_PLATE_GRADE}}）')
        for col in (1, 2, 3):
            _set_cell(table.cell(row, col), '{{IBE_PLATE_MANUFACTURER}}')
        _set_cell(table.cell(row, 4), '{{IBE_PLATE_BATCH}}')
    doc.save(path)


def _prepare_component(path):
    doc = Document(path)
    _prepare_report_common(doc, '{{IBE_COMPONENT_NAME}}出厂检测报告')
    table = doc.tables[0]
    for cell in table.rows[0].cells:
        _set_cell(cell, '规格：{{IBE_ANCHOR_SPEC}}')
    _set_cell(table.cell(2, 2), '{{IBE_ANCHOR_DIAMETER}}±1')
    _set_cell(table.cell(2, 3), '{{IBE_ANCHOR_LENGTH}}±2')
    doc.save(path)


def prepare(source):
    os.makedirs(TEMPLATES_DIR, exist_ok=True)
    outputs = {
        'cover': os.path.join(TEMPLATES_DIR, '01_cover.docx'),
        'toc': os.path.join(TEMPLATES_DIR, '02_toc.docx'),
        'license': os.path.join(TEMPLATES_DIR, '03_business_license.docx'),
        'cert': os.path.join(TEMPLATES_DIR, '04_cert.docx'),
        'plate': os.path.join(TEMPLATES_DIR, '05_plate_report.docx'),
        'component': os.path.join(TEMPLATES_DIR, '06_component_report.docx'),
    }

    _extract(source, outputs['cover'], 0, 26, 0)
    _extract(source, outputs['toc'], 28, 32, 2)
    shared_license = os.path.join(
        ROOT_DIR, 'products', 'isolation_bearing_embedded_parts',
        'templates', '03_business_license.docx')
    if os.path.exists(shared_license):
        shutil.copy2(shared_license, outputs['license'])
    else:
        _extract(source, outputs['license'], 33, 41, 1)
    _extract(source, outputs['cert'], 42, 48, 2)
    _extract(source, outputs['plate'], 49, 53, 2)
    _extract(source, outputs['component'], 67, 78, 2)

    _prepare_cover(outputs['cover'])
    _prepare_cert(outputs['cert'])
    _prepare_plate(outputs['plate'])
    _prepare_component(outputs['component'])
    for path in outputs.values():
        _prune_unused_document_images(path)
        _ensure_numbering_part(path)
    shutil.copy2(outputs['cover'], os.path.join(PRODUCT_DIR, 'template_prepared.docx'))


if __name__ == '__main__':
    if len(sys.argv) != 2:
        raise SystemExit('Usage: prepare_template.py <source-report.docx>')
    prepare(os.path.abspath(sys.argv[1]))
