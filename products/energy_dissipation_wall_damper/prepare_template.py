"""Build templates for energy-dissipation wall-board damper reports.

The source report is a real sample placed in 临时文件.  This script splits it
into reusable section DOCX files and replaces sample values with placeholders.
"""
import copy
import os
import shutil
import zipfile
from pathlib import Path
from posixpath import basename

from docx import Document
from docx.oxml.ns import qn
from lxml import etree


PRODUCT_DIR = Path(__file__).resolve().parent
ROOT_DIR = PRODUCT_DIR.parent.parent
TEMPLATES_DIR = PRODUCT_DIR / 'templates'
REL_NS = 'http://schemas.openxmlformats.org/package/2006/relationships'
OFFICE_REL_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
CONTENT_TYPE_NS = 'http://schemas.openxmlformats.org/package/2006/content-types'
NUMBERING_REL = (
    'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering')
NUMBERING_CONTENT_TYPE = (
    'application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml')
MINIMAL_NUMBERING = b'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'''


def _source_docx():
    candidates = sorted(
        (ROOT_DIR / '临时文件').glob('*.docx'),
        key=lambda p: p.stat().st_size,
        reverse=True,
    )
    for path in candidates:
        if '墙板阻尼器' in path.name:
            return path
    raise FileNotFoundError('未找到“消能减震墙板阻尼器”样本 Word。')


def _extract(source, output, start, end, section_index=0):
    doc = Document(str(source))
    body = doc.element.body
    children = list(body)
    section_index = min(section_index, len(doc.sections) - 1)
    final_section = copy.deepcopy(doc.sections[section_index]._sectPr)
    section_type = final_section.find(qn('w:type'))
    if section_type is not None:
        section_type.set(qn('w:val'), 'continuous')

    for index in range(len(children) - 1, end, -1):
        body.remove(children[index])
    for index in range(start - 1, -1, -1):
        body.remove(children[index])
    for child in list(body):
        nested_section = child.find('.//' + qn('w:sectPr'))
        if nested_section is not None:
            nested_section.getparent().remove(nested_section)
        if child.tag == qn('w:sectPr'):
            body.remove(child)
    body.append(final_section)
    doc.save(str(output))


def _replace_text(doc, old_text, placeholder):
    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, old_text, placeholder)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_text_in_paragraph(paragraph, old_text, placeholder)


def _replace_text_in_paragraph(paragraph, old_text, placeholder):
    runs = paragraph.runs
    if not runs:
        return
    texts = [run.text or '' for run in runs]
    full_text = ''.join(texts)
    idx = full_text.find(old_text)
    if idx < 0:
        return

    end = idx + len(old_text)
    offsets = [0]
    for text in texts:
        offsets.append(offsets[-1] + len(text))

    placed = False
    for run_idx, run in enumerate(runs):
        start_pos, end_pos = offsets[run_idx], offsets[run_idx + 1]
        if end_pos <= idx or start_pos >= end:
            continue
        prefix = texts[run_idx][:max(0, idx - start_pos)]
        suffix = texts[run_idx][max(0, end - start_pos):]
        if not placed:
            run.text = prefix + placeholder + suffix
            placed = True
        else:
            run.text = prefix + suffix


def _replace_value_after_label(doc, label, placeholder):
    for paragraph in doc.paragraphs:
        runs = paragraph.runs
        if not runs:
            continue
        full_text = ''.join(run.text or '' for run in runs)
        idx = full_text.find(label)
        if idx < 0:
            continue
        value = full_text[idx + len(label):].strip()
        if value:
            _replace_text_in_paragraph(paragraph, value, placeholder)


def _ensure_numbering_part(path):
    with zipfile.ZipFile(path, 'r') as source:
        files = {name: source.read(name) for name in source.namelist()}

    files.setdefault('word/numbering.xml', MINIMAL_NUMBERING)

    rel_name = 'word/_rels/document.xml.rels'
    rels = etree.fromstring(files[rel_name])
    has_numbering_rel = any(
        rel.get('Type') == NUMBERING_REL
        for rel in rels.findall(f'{{{REL_NS}}}Relationship')
    )
    if not has_numbering_rel:
        used = []
        for rel in rels.findall(f'{{{REL_NS}}}Relationship'):
            rid = rel.get('Id', '')
            if rid.startswith('rId') and rid[3:].isdigit():
                used.append(int(rid[3:]))
        new_rel = etree.Element(f'{{{REL_NS}}}Relationship')
        new_rel.set('Id', f'rId{(max(used) if used else 0) + 1}')
        new_rel.set('Type', NUMBERING_REL)
        new_rel.set('Target', 'numbering.xml')
        rels.append(new_rel)
        files[rel_name] = etree.tostring(
            rels, xml_declaration=True, encoding='UTF-8', standalone=True)

    content_types = etree.fromstring(files['[Content_Types].xml'])
    has_numbering_content_type = any(
        node.get('PartName') == '/word/numbering.xml'
        for node in content_types.findall(f'{{{CONTENT_TYPE_NS}}}Override')
    )
    if not has_numbering_content_type:
        override = etree.Element(f'{{{CONTENT_TYPE_NS}}}Override')
        override.set('PartName', '/word/numbering.xml')
        override.set('ContentType', NUMBERING_CONTENT_TYPE)
        content_types.append(override)
        files['[Content_Types].xml'] = etree.tostring(
            content_types, xml_declaration=True, encoding='UTF-8', standalone=True)

    with zipfile.ZipFile(path, 'w', zipfile.ZIP_DEFLATED) as target:
        for name, content in files.items():
            target.writestr(name, content)


def _prune_unused_document_images(path):
    with zipfile.ZipFile(path, 'r') as source:
        files = {name: source.read(name) for name in source.namelist()}

    document = etree.fromstring(files['word/document.xml'])
    used_ids = set(document.xpath(
        '//@r:embed | //@r:id', namespaces={'r': OFFICE_REL_NS}))
    rel_name = 'word/_rels/document.xml.rels'
    rels = etree.fromstring(files[rel_name])
    for rel in list(rels):
        if rel.get('Type', '').endswith('/image') and rel.get('Id') not in used_ids:
            rels.remove(rel)
    files[rel_name] = etree.tostring(
        rels, xml_declaration=True, encoding='UTF-8', standalone=True)

    referenced_media = set()
    for name, content in files.items():
        if not name.endswith('.rels'):
            continue
        root = etree.fromstring(content)
        for rel in root.findall(f'{{{REL_NS}}}Relationship'):
            target = rel.get('Target', '')
            if 'media/' in target:
                referenced_media.add(basename(target))
    for name in list(files):
        if name.startswith('word/media/') and basename(name) not in referenced_media:
            del files[name]

    with zipfile.ZipFile(path, 'w', zipfile.ZIP_DEFLATED) as target:
        for name, content in files.items():
            target.writestr(name, content)


def _save_prepared(path, replacements=None, labels=None):
    doc = Document(str(path))
    for label, placeholder in (labels or {}).items():
        _replace_value_after_label(doc, label, placeholder)
    for old_text, placeholder in (replacements or {}).items():
        _replace_text(doc, old_text, placeholder)
    doc.save(str(path))
    _ensure_numbering_part(path)
    _prune_unused_document_images(path)


def build():
    source = _source_docx()
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)

    section_defs = [
        ('01_cover.docx', 0, 15),
        ('02_toc.docx', 16, 22),
        ('03_business_license.docx', 23, 26),
        ('04_cert.docx', 27, 31),
        ('05_inspection_report.docx', 42, 56, 4),
        ('06_material_title.docx', 77, 77),
    ]
    for item in section_defs:
        name, start, end = item[:3]
        section_index = item[3] if len(item) > 3 else 0
        _extract(source, TEMPLATES_DIR / name, start, end, section_index)

    _save_prepared(
        TEMPLATES_DIR / '01_cover.docx',
        labels={'项目名称：': '{{WALL_PROJECT_NAME}}'},
        replacements={
            '四川融海运通抗震科技有限责任公司': '{{WALL_SUPPLIER}}',
            '四川省成都市新津区普兴街道清云南路256号': '{{WALL_ADDRESS}}',
            '18180838170': '{{WALL_PHONE}}',
            '2026年6月': '{{WALL_REPORT_DATE}}',
        },
    )
    _save_prepared(TEMPLATES_DIR / '02_toc.docx')
    _save_prepared(TEMPLATES_DIR / '03_business_license.docx')
    _save_prepared(
        TEMPLATES_DIR / '04_cert.docx',
        replacements={
            '消能减震墙板阻尼器': '{{WALL_PRODUCT_NAME}}',
            'QB-1800-3360': '{{WALL_MODEL}}',
            '2026年6月': '{{WALL_PRODUCTION_DATE}}',
            '01': '{{WALL_INSPECTOR}}',
            '6件': '{{WALL_QTY_WITH_UNIT}}',
        },
    )
    _save_prepared(
        TEMPLATES_DIR / '05_inspection_report.docx',
        replacements={
            '消能减震墙板阻尼器': '{{WALL_PRODUCT_NAME}}',
            'JG/T209-2012': '{{WALL_STANDARD}}',
            'QB-1800-3360': '{{WALL_MODEL}}',
            'C35': '{{WALL_STRENGTH_GRADE}}',
            '合格': '{{WALL_PASS}}',
        },
    )
    _save_prepared(TEMPLATES_DIR / '06_material_title.docx')

    shutil.copy2(TEMPLATES_DIR / '03_business_license.docx',
                 PRODUCT_DIR / 'latest_business_license.docx')


if __name__ == '__main__':
    build()
