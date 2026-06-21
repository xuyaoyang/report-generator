"""
Prepare the isolation bearing template by scanning for yellow-highlighted
runs and replacing them with {{FIELD_XXX}} placeholders.

Unlike the VD template (which had no highlights), the IB template was
originally prepared by manually highlighting variable fields in yellow.

Produces: template_prepared.docx + param_mapping.json
"""
import json, os, re
from docx import Document
from docx.oxml.ns import qn
from lxml import etree


def _is_highlighted(run_elem):
    """Check if a run element has yellow highlight."""
    rPr = run_elem.find(qn('w:rPr'))
    if rPr is None:
        return False
    hl = rPr.find(qn('w:highlight'))
    if hl is None:
        return False
    return hl.get(qn('w:val')) == 'yellow'


def _process_container(container, fields, field_counter, location_prefix):
    """Find and replace all yellow-highlighted text in a container (paragraph or cell).

    Groups consecutive highlighted runs into a single field.
    Each non-consecutive highlighted segment gets its own field ID.
    location_prefix: e.g. 'paragraph[13]' or 'table[0].cell[1,2]'
    """
    paras = container.paragraphs if hasattr(container, 'paragraphs') else [container]

    for para in paras:
        runs = para.runs
        if not runs:
            continue

        # Find groups of consecutive highlighted runs
        groups = []
        current_group = []
        for i, run in enumerate(runs):
            if _is_highlighted(run._element):
                current_group.append(i)
            else:
                if current_group:
                    groups.append(current_group)
                    current_group = []
        if current_group:
            groups.append(current_group)

        if not groups:
            continue

        # Process each group (in reverse to preserve indices)
        for group in reversed(groups):
            # Collect original text from highlighted runs
            original_parts = []
            for idx in group:
                original_parts.append(runs[idx].text or '')

            original = ''.join(original_parts)
            if not original.strip():
                continue

            field_counter[0] += 1
            placeholder = f'{{{{FIELD_{field_counter[0]:03d}}}}}'

            # Determine human-readable label from context
            label = _infer_label(location_prefix, original)

            fields.append({
                'id': field_counter[0],
                'placeholder': placeholder,
                'original': original,
                'location': f'{location_prefix} ({label})' if label else location_prefix,
            })

            # Put placeholder in first highlighted run, clear others
            runs[group[0]].text = placeholder
            for idx in group[1:]:
                runs[idx].text = ''


def _infer_label(location, text):
    """Try to infer a human-readable label for the field."""
    text = text.strip()

    # Common patterns
    if '项目名称' in location or (len(text) > 10 and '项目' in text):
        return '项目名称'
    if '供应商' in location or '融海运通' in text:
        return '产品供应商'
    if '地址' in location:
        return '制造地址'
    if re.match(r'^\d{10,11}$', text):
        return '联系电话'
    if re.match(r'^\d{4}\s*年\s*\d{1,2}\s*月', text):
        return '制造日期'

    return ''  # let categorization handle it later


def categorize(fields):
    """Assign each field to a category based on its location.

    For tables, use table index ranges (determined by document structure).
    For paragraphs, use label-based matching.
    """
    categories = {
        'cover_info': [],
        'certificates': [],
        'visual': [],
        'mechanical': [],
        'other': [],
    }

    for f in fields:
        loc = f['location']

        if 'paragraph[' in loc:
            # All paragraph fields are cover info (project info, inspector, etc.)
            categories['cover_info'].append(f)

        elif 'table[' in loc:
            t_num = int(re.search(r'table\[(\d+)\]', loc).group(1))

            if 0 <= t_num <= 5:
                # Certificate tables
                categories['certificates'].append(f)
            elif 6 <= t_num <= 8:
                # Mechanical tables (Table 6 = summary, Tables 7-8 = detail)
                categories['mechanical'].append(f)
            elif 9 <= t_num <= 14:
                # Visual inspection tables
                categories['visual'].append(f)
            else:
                categories['other'].append(f)
        else:
            categories['other'].append(f)

    return categories


def prepare_template(template_path, output_path, mapping_path):
    doc = Document(template_path)
    fields = []
    field_counter = [0]

    # 1. Process cover paragraphs (paragraphs before first table)
    for p_idx, para in enumerate(doc.paragraphs):
        _process_container(para, fields, field_counter, f'paragraph[{p_idx}]')

    # 2. Process all tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                _process_container(cell, fields, field_counter,
                                   f'table[{t_idx}].cell[{r_idx},{c_idx}]')

    # 3. Remove yellow highlighting from all remaining runs (clean up)
    for para in doc.paragraphs:
        for run in para.runs:
            rPr = run._element.find(qn('w:rPr'))
            if rPr is not None:
                hl = rPr.find(qn('w:highlight'))
                if hl is not None:
                    rPr.remove(hl)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        rPr = run._element.find(qn('w:rPr'))
                        if rPr is not None:
                            hl = rPr.find(qn('w:highlight'))
                            if hl is not None:
                                rPr.remove(hl)

    # Save
    doc.save(output_path)

    # Categorize and save mapping
    categories = categorize(fields)
    mapping = {
        'total_fields': len(fields),
        'categories': {
            k: [{'placeholder': f['placeholder'],
                 'original': f['original'],
                 'location': f['location']} for f in v]
            for k, v in categories.items()
        },
    }

    with open(mapping_path, 'w', encoding='utf-8') as f:
        json.dump(mapping, f, ensure_ascii=False, indent=2)

    print(f'Total fields: {len(fields)}')
    for cat, items in mapping['categories'].items():
        print(f'  {cat}: {len(items)} fields')

    return mapping


if __name__ == '__main__':
    base = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(base, 'template.docx')
    output_path = os.path.join(base, 'template_prepared.docx')
    mapping_path = os.path.join(base, 'param_mapping.json')

    prepare_template(template_path, output_path, mapping_path)
    print(f'\nSaved: {output_path}')
    print(f'Saved: {mapping_path}')
