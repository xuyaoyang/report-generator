"""Build section templates from an isolation-bearing embedded-parts report."""
import copy
import os
import shutil
import sys
import tempfile
import zipfile
from posixpath import basename

from lxml import etree

from docx import Document
from docx.oxml.ns import qn


PRODUCT_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATES_DIR = os.path.join(PRODUCT_DIR, 'templates')
REL_NS = 'http://schemas.openxmlformats.org/package/2006/relationships'
OFFICE_REL_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'


def _extract(source, output, start, end, section_index):
    doc = Document(source)
    body = doc.element.body
    children = list(body)
    final_section = copy.deepcopy(doc.sections[section_index]._sectPr)
    section_type = final_section.find(qn('w:type'))
    if section_type is None:
        section_type = final_section.makeelement(qn('w:type'))
        final_section.insert(0, section_type)
    section_type.set(qn('w:val'), 'continuous')
    for index in range(len(children) - 1, end, -1):
        body.remove(children[index])
    for index in range(start - 1, -1, -1):
        body.remove(children[index])
    for child in list(body):
        nested_section = child.find('.//' + qn('w:sectPr'))
        if nested_section is not None:
            nested_section.getparent().remove(nested_section)
        if child.tag == qn('w:sectPr'):
            body.remove(child)
    body.append(final_section)
    doc.save(output)


def _prune_unused_document_images(path):
    with zipfile.ZipFile(path, 'r') as source:
        files = {name: source.read(name) for name in source.namelist()}

    document = etree.fromstring(files['word/document.xml'])
    used_ids = set(document.xpath(
        '//@r:embed | //@r:id', namespaces={'r': OFFICE_REL_NS}))
    rel_name = 'word/_rels/document.xml.rels'
    rels = etree.fromstring(files[rel_name])
    for rel in list(rels):
        if rel.get('Type', '').endswith('/image') and rel.get('Id') not in used_ids:
            rels.remove(rel)
    files[rel_name] = etree.tostring(
        rels, xml_declaration=True, encoding='UTF-8', standalone=True)

    referenced_media = set()
    for name, content in files.items():
        if not name.endswith('.rels'):
            continue
        root = etree.fromstring(content)
        for rel in root.findall(f'{{{REL_NS}}}Relationship'):
            target = rel.get('Target', '')
            if 'media/' in target:
                referenced_media.add(basename(target))
    for name in list(files):
        if name.startswith('word/media/') and basename(name) not in referenced_media:
            del files[name]

    fd, temp_path = tempfile.mkstemp(suffix='.docx', dir=os.path.dirname(path))
    os.close(fd)
    try:
        with zipfile.ZipFile(temp_path, 'w', zipfile.ZIP_DEFLATED) as target:
            for name, content in files.items():
                target.writestr(name, content)
        os.replace(temp_path, path)
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def _set_paragraph(paragraph, value):
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ''
    else:
        paragraph.add_run(value)


def _set_cell(cell, value):
    paragraph = cell.paragraphs[0]
    _set_paragraph(paragraph, value)
    for extra in cell.paragraphs[1:]:
        _set_paragraph(extra, '')


def _replace_prefixed_paragraph(doc, prefix, value):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            _set_paragraph(paragraph, value)
            return
    raise ValueError(f'Paragraph not found: {prefix}')


def _prepare_cover(path):
    doc = Document(path)
    _replace_prefixed_paragraph(doc, '\u56db\u5ddd\u878d\u6d77', '{{IBE_SUPPLIER}}')
    _replace_prefixed_paragraph(doc, '\u9884\u57cb\u4ef6\u51fa\u5382\u8d44\u6599', '{{IBE_TITLE}}')
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith('\u5de5\u7a0b\u540d\u79f0'):
            paragraph.runs[2].text = '{{IBE_PROJECT}}'
            break
    _replace_prefixed_paragraph(doc, '\u4ea7\u54c1\u4f9b\u5e94\u5546', '\u4ea7\u54c1\u4f9b\u5e94\u5546\uff1a{{IBE_SUPPLIER}}')
    _replace_prefixed_paragraph(doc, '\u751f\u4ea7\u5730\u5740', '\u751f\u4ea7\u5730\u5740\uff1a{{IBE_ADDRESS}}')
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith('\u7535'):
            paragraph.runs[-1].text = '{{IBE_PHONE}}'
        if paragraph.text.strip().startswith('\u62a5\u544a\u65e5\u671f'):
            paragraph.runs[1].text = '{{IBE_DATE}}'
            for run in paragraph.runs[2:]:
                run.text = ''
    doc.save(path)


def _prepare_cert(path):
    doc = Document(path)
    table = doc.tables[0]
    for index in range(1, 14):
        values = ('NAME', 'SPEC', 'BATCH', 'DATE', 'QTY')
        for col, field in enumerate(values):
            _set_cell(table.cell(index, col), f'{{{{IBE_CERT_{index}_{field}}}}}')
    doc.save(path)


def _prepare_anchor(path):
    doc = Document(path)
    _replace_prefixed_paragraph(doc, '\u56db\u5ddd\u878d\u6d77', '{{IBE_SUPPLIER}}')
    table = doc.tables[0]
    _set_cell(table.cell(0, 3), '{{IBE_ANCHOR_SPEC}}')
    _set_cell(table.cell(3, 5), '{{IBE_ANCHOR_DIAMETER}}')
    _set_cell(table.cell(3, 6), '{{IBE_ANCHOR_LENGTH}}')
    _set_cell(table.cell(7, 5), '{{IBE_SLEEVE_OD}}')
    _set_cell(table.cell(7, 6), '{{IBE_SLEEVE_LENGTH}}')
    for row in range(11, 18):
        _set_cell(table.cell(row, 0), '{{IBE_ANCHOR_GRADE}}')
        _set_cell(table.cell(row, 1), '{{IBE_ANCHOR_MANUFACTURER}}')
        _set_cell(table.cell(row, 2), '{{IBE_ANCHOR_BATCH}}')
    for row in range(18, 21):
        _set_cell(table.cell(row, 1), '{{IBE_SUPPLIER}}')
        _set_cell(table.cell(row, 2), '{{IBE_ANCHOR_BATCH}}')
    doc.save(path)


def _prepare_plate(path):
    doc = Document(path)
    _replace_prefixed_paragraph(doc, '\u56db\u5ddd\u878d\u6d77', '{{IBE_SUPPLIER}}')
    table = doc.tables[0]
    _set_cell(table.cell(0, 5), '{{IBE_PLATE_SPEC}}')
    _set_cell(table.cell(3, 3), '{{IBE_PLATE_SIDE_STANDARD}}')
    _set_cell(table.cell(3, 6), '{{IBE_PLATE_THICKNESS_STANDARD}}')
    for row in range(7, 10):
        _set_cell(table.cell(row, 0), '{{IBE_PLATE_GRADE}}')
        _set_cell(table.cell(row, 1), '{{IBE_PLATE_MANUFACTURER}}')
        _set_cell(table.cell(row, 4), '{{IBE_PLATE_BATCH}}')
    doc.save(path)


def prepare(source):
    os.makedirs(TEMPLATES_DIR, exist_ok=True)

    cover = os.path.join(TEMPLATES_DIR, '01_cover.docx')
    toc = os.path.join(TEMPLATES_DIR, '02_toc.docx')
    license_path = os.path.join(TEMPLATES_DIR, '03_business_license.docx')
    cert = os.path.join(TEMPLATES_DIR, '04_cert.docx')
    anchor = os.path.join(TEMPLATES_DIR, '05_anchor_report.docx')
    plate = os.path.join(TEMPLATES_DIR, '06_plate_report.docx')

    _extract(source, cover, 0, 26, 0)
    _extract(source, toc, 28, 37, 2)
    _extract(source, license_path, 27, 27, 1)
    _extract(source, cert, 43, 50, 2)
    _extract(source, anchor, 51, 58, 2)
    _extract(source, plate, 72, 79, 2)

    _prepare_cover(cover)
    _prepare_cert(cert)
    _prepare_anchor(anchor)
    _prepare_plate(plate)
    for path in (cover, toc, license_path, cert, anchor, plate):
        _prune_unused_document_images(path)
    shutil.copy2(cover, os.path.join(PRODUCT_DIR, 'template_prepared.docx'))


if __name__ == '__main__':
    if len(sys.argv) != 2:
        raise SystemExit('Usage: prepare_template.py <source-report.docx>')
    prepare(os.path.abspath(sys.argv[1]))
