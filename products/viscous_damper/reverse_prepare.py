"""
Reverse prepare: convert a generated report back to template_prepared.docx
and template.docx by replacing filled values with {{FIELD_XXX}} placeholders.

Uses the current template_prepared.docx as a "mask" to identify where labels
end and values begin — far more robust than matching against original text.

Workflow for adjusting template formatting:
  1. Generate a report with sample data
  2. Open the report in Word, adjust formatting (fonts, table widths, etc.)
  3. Run: python reverse_prepare.py <adjusted_report.docx>
  4. This produces updated template.docx and template_prepared.docx
     with the adjusted formatting and restored placeholders.

Limitation: the report must have the SAME paragraph/table structure as the
template (no dynamic adjustment). For VD, generate with exactly 1 model to
match the template structure.
"""
import json
import os
import re
import shutil
import sys
from docx import Document


def _put_text_in_first_overlapping_run(container, new_text):
    """Replace all text in a paragraph or cell with new_text,
    putting it in the first run and clearing others."""
    for para in (container.paragraphs if hasattr(container, 'paragraphs') else [container]):
        runs = para.runs
        if not runs:
            continue
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ''


def reverse_prepare(report_path, prepared_template_path, mapping_path,
                    output_template_path, output_prepared_path):
    """
    Convert a generated report back to template files.

    Uses the prepared template as a mask: for each placeholder location,
    reads the label+suffix from the prepared template, finds them in the
    report, and replaces the value between them with the placeholder.
    """
    with open(mapping_path, 'r', encoding='utf-8') as f:
        mapping = json.load(f)

    report_doc = Document(report_path)
    mask_doc = Document(prepared_template_path)

    # Collect all fields from all categories
    all_fields = []
    for cat, items in mapping['categories'].items():
        for item in items:
            all_fields.append(item)

    replaced = 0
    failed = 0

    for field in all_fields:
        ph = field['placeholder']
        loc = field['location']

        para_match = re.search(r'paragraph\[(\d+)\]', loc)
        table_match = re.search(r'table\[(\d+)\]\.cell\[(\d+),(\d+)\]', loc)

        if para_match:
            p_idx = int(para_match.group(1))
            if p_idx >= len(report_doc.paragraphs):
                failed += 1
                continue

            report_para = report_doc.paragraphs[p_idx]
            report_text = report_para.text

            if ph in report_text:
                continue  # already has placeholder

            # Get mask text from prepared template
            if p_idx < len(mask_doc.paragraphs):
                mask_text = mask_doc.paragraphs[p_idx].text
            else:
                mask_text = ''

            # Split mask by this placeholder to get label + suffix
            if ph in mask_text:
                label, suffix = mask_text.split(ph, 1)
            else:
                # Fallback: use original text from mapping
                original = field.get('original', '')
                label = ''
                suffix = ''
                if original and original in mask_text:
                    label = mask_text.split(original, 1)[0]

            # Find label and suffix in report text, replace middle with placeholder
            if label and label in report_text:
                after_label = report_text.split(label, 1)[1]
                if suffix and suffix in after_label:
                    # Precise: replace just the value between label and suffix
                    value_part = after_label.split(suffix, 1)[0]
                    new_text = report_text.replace(
                        label + value_part + suffix,
                        label + ph + suffix, 1
                    )
                else:
                    # No suffix: replace everything after label
                    new_text = label + ph
                _put_text_in_first_overlapping_run(report_para, new_text)
                replaced += 1
            else:
                # Label not found — fall back to original-text matching
                original = field.get('original', '')
                if original and original in report_text:
                    new_text = report_text.replace(original, ph, 1)
                    _put_text_in_first_overlapping_run(report_para, new_text)
                    replaced += 1
                else:
                    print(f'  WARN {ph}: cannot find label or original in '
                          f'"{report_text[:60]}..."')
                    failed += 1

        elif table_match:
            t_idx = int(table_match.group(1))
            row = int(table_match.group(2))
            col = int(table_match.group(3))

            if t_idx >= len(report_doc.tables):
                failed += 1
                continue

            table = report_doc.tables[t_idx]
            if row >= len(table.rows) or col >= len(table.rows[row].cells):
                failed += 1
                continue

            cell = table.cell(row, col)
            cell_text = cell.text

            if ph in cell_text:
                continue

            # Get mask text from prepared template
            mask_text = ''
            if t_idx < len(mask_doc.tables):
                mask_table = mask_doc.tables[t_idx]
                if row < len(mask_table.rows) and col < len(mask_table.rows[row].cells):
                    mask_text = mask_table.cell(row, col).text

            # Try mask-based replacement
            if ph in mask_text:
                label, suffix = mask_text.split(ph, 1)
            else:
                label = ''
                suffix = ''

            if label and label in cell_text:
                after_label = cell_text.split(label, 1)[1]
                if suffix and suffix in after_label:
                    value_part = after_label.split(suffix, 1)[0]
                    new_text = cell_text.replace(
                        label + value_part + suffix,
                        label + ph + suffix, 1
                    )
                else:
                    new_text = label + ph
                _put_text_in_first_overlapping_run(cell, new_text)
                replaced += 1
            else:
                # Fallback: original text matching
                original = field.get('original', '')
                if original and original in cell_text:
                    new_text = cell_text.replace(original, ph, 1)
                    _put_text_in_first_overlapping_run(cell, new_text)
                    replaced += 1
                elif cell_text.strip():
                    # Cell has only the value (no label)
                    _put_text_in_first_overlapping_run(cell, ph)
                    replaced += 1
                else:
                    failed += 1

    # Save both template files
    report_doc.save(output_prepared_path)
    shutil.copy2(output_prepared_path, output_template_path)

    print(f'\nReplaced: {replaced}, Failed: {failed}')
    print(f'Template: {output_template_path}')
    print(f'Prepared: {output_prepared_path}')
    return replaced, failed


if __name__ == '__main__':
    base = os.path.dirname(os.path.abspath(__file__))
    mapping_path = os.path.join(base, 'param_mapping.json')
    prepared_template = os.path.join(base, 'template_prepared.docx')
    output_template = os.path.join(base, 'template.docx')
    output_prepared = os.path.join(base, 'template_prepared.docx')

    if len(sys.argv) < 2:
        print('Usage: python reverse_prepare.py <generated_report.docx>')
        print()
        print('Converts a generated report back to template files by')
        print('replacing filled values with {{FIELD_XXX}} placeholders.')
        print()
        print('Workflow:')
        print('  1. Generate a report')
        print('  2. Adjust formatting in Word')
        print('  3. Run this script on the adjusted report')
        print('  4. Template files are updated with new formatting')
        sys.exit(1)

    report_path = sys.argv[1]
    if not os.path.exists(report_path):
        print(f'Error: file not found: {report_path}')
        sys.exit(1)

    # Back up current templates
    for fname in ['template.docx', 'template_prepared.docx']:
        src = os.path.join(base, fname)
        if os.path.exists(src):
            bak = src + '.bak'
            shutil.copy2(src, bak)
            print(f'Backed up: {fname} -> {fname}.bak')

    reverse_prepare(report_path, prepared_template, mapping_path,
                    output_template, output_prepared)
