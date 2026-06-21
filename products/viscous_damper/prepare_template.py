"""
Prepare the viscous damper template by replacing sample data with
{{FIELD_XXX}} placeholders. Unlike the isolation bearing template,
this one has NO yellow highlights, so we use position + content
pattern matching to identify variable fields.

Produces: template_prepared.docx + param_mapping.json
"""
import json
import os
import re
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn


def _replace_text_in_runs(container, old_text, placeholder):
    """Replace old_text with placeholder across runs, preserving per-run formatting.

    Unlike the naive approach of stuffing everything into the first run,
    this only modifies runs that intersect the matched text range. Runs
    before/after the match keep their original text and formatting.
    """
    made_change = False
    for para in (container.paragraphs if hasattr(container, 'paragraphs') else [container]):
        runs = para.runs
        if not runs:
            continue

        run_texts = [r.text or '' for r in runs]
        full_text = ''.join(run_texts)

        idx = full_text.find(old_text)
        if idx == -1:
            continue

        made_change = True
        old_end = idx + len(old_text)
        placeholder_placed = False

        # Build cumulative character offsets for each run
        offsets = [0]
        for t in run_texts:
            offsets.append(offsets[-1] + len(t))

        for i, run in enumerate(runs):
            s, e = offsets[i], offsets[i + 1]
            if e <= idx or s >= old_end:
                continue  # no overlap, keep as-is

            prefix = run_texts[i][:max(0, idx - s)]
            suffix = run_texts[i][max(0, old_end - s):]

            if not placeholder_placed:
                run.text = prefix + placeholder + suffix
                placeholder_placed = True
            else:
                run.text = prefix + suffix

    return made_change


def _replace_value_after_label(container, label, placeholder):
    """Replace text after a known label with placeholder.

    E.g., for paragraph text "项目名称：某某项目" with label "项目名称：",
    replaces "某某项目" with placeholder, yielding "项目名称：{{FIELD_001}}".

    The label text is retained — only the value portion after it is replaced.
    Works regardless of what the current value text is.
    """
    for para in (container.paragraphs if hasattr(container, 'paragraphs') else [container]):
        runs = para.runs
        if not runs:
            continue

        run_texts = [r.text or '' for r in runs]
        full_text = ''.join(run_texts)

        label_idx = full_text.find(label)
        if label_idx == -1:
            return False

        # Value starts right after the label, ends at end of text
        value_start = label_idx + len(label)
        value_end = len(full_text)

        if value_start >= value_end:
            # No value text — just append placeholder
            if runs:
                runs[-1].text = (runs[-1].text or '') + placeholder
            return True

        # Replace the value portion [value_start, value_end) with placeholder
        offsets = [0]
        for t in run_texts:
            offsets.append(offsets[-1] + len(t))

        placeholder_placed = False
        for i, run in enumerate(runs):
            s, e = offsets[i], offsets[i + 1]
            if e <= value_start:
                continue  # before value
            if s >= value_end:
                continue  # after value

            prefix = run_texts[i][:max(0, value_start - s)]
            suffix = run_texts[i][max(0, value_end - s):]

            if not placeholder_placed:
                run.text = prefix + placeholder + suffix
                placeholder_placed = True
            else:
                run.text = prefix + suffix

        return True


def _replace_all_text_in_container(container, placeholder):
    """Replace ALL text in a paragraph or cell with the placeholder.

    Puts the placeholder in the first run, clears other runs.
    Useful for table cells whose entire content is the field value.
    """
    for para in (container.paragraphs if hasattr(container, 'paragraphs') else [container]):
        runs = para.runs
        if not runs:
            continue
        runs[0].text = placeholder
        for run in runs[1:]:
            run.text = ''


def prepare_template(template_path, output_path, mapping_path):
    doc = Document(template_path)
    fields = []
    field_id = 0

    def add_field(location, original_text):
        nonlocal field_id
        field_id += 1
        placeholder = f'{{{{FIELD_{field_id:03d}}}}}'
        fields.append({
            'id': field_id,
            'placeholder': placeholder,
            'original': original_text,
            'location': location,
        })
        return placeholder

    # ============================================================
    # 1. COVER PARAGRAPHS — use label-based replacement so the user
    #    can freely edit the value text (project name, company, etc.)
    #    without breaking the prepare script.
    # ============================================================
    cover_specs = [
        (13, '项目名称：', '项目名称'),
        (29, '产品供应商：', '产品供应商'),
        (30, '生产地址：', '生产地址'),
        (31, '电    话：', '电话'),
        (32, '报告日期：', '报告日期'),
    ]
    for p_idx, label, field_name in cover_specs:
        para = doc.paragraphs[p_idx]
        full_text = para.text
        after_label = full_text.split(label, 1)[1] if label in full_text else full_text
        old_val = after_label.strip()
        ph = add_field(f'paragraph[{p_idx}] ({field_name})', old_val)
        if label in full_text:
            _replace_value_after_label(para, label, ph)
        else:
            _replace_all_text_in_container(para, ph)

    # ============================================================
    # 2. CERTIFICATE TABLES (Table 0-7, each 4x4)
    # ============================================================
    # Fields per cert: model, date, serial, inspector, standard, dept, verdict
    cert_field_specs = [
        (0, 3, 'VFD-NL×350×45', '规格型号'),
        (1, 1, '2024 年 12 月', '生产日期'),
        (1, 3, None, '出厂编号'),  # serial varies per cert, handled below
        (2, 1, '01', '检验员'),
        (2, 3, 'JG/T209-2012', '检验依据'),
        (3, 1, '质量部', '检验部门'),
        (3, 3, '合格', '检验结果'),
    ]

    cert_serials = [
        'RB124014855', 'RB124014856', 'RB124014857', 'RB124014858',
        'RB124014859', 'RB124014860', 'RB124014861', 'RB124014862',
    ]

    for t_idx in range(8):
        table = doc.tables[t_idx]
        for row_idx, col_idx, old_val, label in cert_field_specs:
            cell = table.cell(row_idx, col_idx)
            actual_val = cert_serials[t_idx] if old_val is None else old_val
            ph = add_field(
                f'table[{t_idx}].cell[{row_idx},{col_idx}] ({label})',
                actual_val
            )
            _replace_all_text_in_container(cell, ph)

    # ============================================================
    # 3. APPEARANCE TEST REPORT HEADER PARAGRAPHS
    # ============================================================
    # P[0136] company name — entire paragraph is the field value
    para136 = doc.paragraphs[136]
    old_val_136 = para136.text.strip()
    ph = add_field('paragraph[136] (外观报告-公司名称)', old_val_136)
    _replace_all_text_in_container(para136, ph)

    # P[0139] inspector / reviewer / date
    para139 = doc.paragraphs[139]
    replacements_139 = [
        ('何永琴', '外观报告-检验员'),
        ('童俊豪', '外观报告-审核'),
        ('2024 年 12 月', '外观报告-检测日期'),
    ]
    for old_val, label in replacements_139:
        ph = add_field(f'paragraph[139] ({label})', old_val)
        _replace_text_in_runs(para139, old_val, ph)

    # ============================================================
    # 4. APPEARANCE TEST TABLE (Table 8, 12x6 with merged cells)
    # ============================================================
    # All cells below are value-only (labels are in separate header cells),
    # so we use _replace_all_text_in_container for robustness.
    t8 = doc.tables[8]

    # Row 0, C1: project name (value-only, label in header row)
    cell = t8.cell(0, 1)
    old_val = cell.text.strip()
    ph = add_field('table[8].cell[0,1] (外观-项目名称)', old_val)
    _replace_all_text_in_container(cell, ph)

    # Row 2, C3: product model
    cell = t8.cell(2, 3)
    old_val = cell.text.strip()
    ph = add_field('table[8].cell[2,3] (外观-产品型号)', old_val)
    _replace_all_text_in_container(cell, ph)

    # Rows 4-10: test data (cols 3=实测值, 4=数量, 5=结论)
    test_items_8 = [
        (4, 3, '表面平滑度-实测值'),
        (4, 4, '表面平滑度-数量'),
        (4, 5, '表面平滑度-结论'),
        (5, 3, '机械损伤-实测值'),
        (5, 4, '机械损伤-数量'),
        (5, 5, '机械损伤-结论'),
        (6, 3, '锈蚀毛刺-实测值'),
        (6, 4, '锈蚀毛刺-数量'),
        (6, 5, '锈蚀毛刺-结论'),
        (7, 3, '无渗漏-实测值'),
        (7, 4, '无渗漏-数量'),
        (7, 5, '无渗漏-结论'),
        (8, 3, '产品标识-实测值'),
        (8, 4, '产品标识-数量'),
        (8, 5, '产品标识-结论'),
        (9, 3, '长度偏差-实测值'),
        (9, 4, '长度偏差-数量'),
        (9, 5, '长度偏差-结论'),
        (10, 3, '截面尺寸偏差-实测值'),
        (10, 4, '截面尺寸偏差-数量'),
        (10, 5, '截面尺寸偏差-结论'),
    ]
    for row, col, label in test_items_8:
        cell = t8.cell(row, col)
        old_val = cell.text.strip()
        ph = add_field(f'table[8].cell[{row},{col}] ({label})', old_val)
        _replace_all_text_in_container(cell, ph)

    # Row 11, C0: conclusion text with model name embedded (multi-field)
    cell = t8.cell(11, 0)
    ph = add_field('table[8].cell[11,0] (外观-检验结论-型号)',
                   'VFD-NL×350×45')
    _replace_text_in_runs(cell, 'VFD-NL×350×45', ph)

    ph = add_field('table[8].cell[11,0] (外观-检验结论-项目名称)',
                   '该项目黏滞阻尼器')
    _replace_text_in_runs(cell, '该项目黏滞阻尼器', ph)

    # Row 11, C1: inspection department (value-only)
    cell = t8.cell(11, 1)
    old_val = cell.text.strip()
    ph = add_field('table[8].cell[11,1] (外观-检验部门)', old_val)
    _replace_all_text_in_container(cell, ph)

    # ============================================================
    # 5. MECHANICAL TEST TABLE (Table 9, 14x6 with merged cells)
    # ============================================================
    # All cells below are value-only (labels are in separate header cells).
    t9 = doc.tables[9]

    # Row 0, C1: project name
    cell = t9.cell(0, 1)
    old_val = cell.text.strip()
    ph = add_field('table[9].cell[0,1] (力学-项目名称)', old_val)
    _replace_all_text_in_container(cell, ph)

    # Row 1, C1: model
    cell = t9.cell(1, 1)
    old_val = cell.text.strip()
    ph = add_field('table[9].cell[1,1] (力学-规格型号)', old_val)
    _replace_all_text_in_container(cell, ph)

    # Row 4: design parameter values (6 cols)
    design_params = [
        (0, '最大阻尼力(kN)'),
        (1, '设计位移(mm)'),
        (2, '极限位移(mm)'),
        (3, '阻尼系数'),
        (4, '阻尼力指数'),
        (5, '结构基频(Hz)'),
    ]
    for col, label in design_params:
        cell = t9.cell(4, col)
        old_val = cell.text.strip()
        ph = add_field(f'table[9].cell[4,{col}] (力学-{label})', old_val)
        _replace_all_text_in_container(cell, ph)

    # Rows 7-10: test results (col 3 = result, col 4 = verdict)
    mech_results = [
        (7, 3, '设计阻尼力-检测结果'),
        (7, 4, '设计阻尼力-判定'),
        (8, 3, '极限位移-检测结果'),
        (8, 4, '极限位移-判定'),
        (9, 3, '阻尼系数-检测结果'),
        (9, 4, '阻尼系数-判定'),
        (10, 3, '阻尼指数-检测结果'),
        (10, 4, '阻尼指数-判定'),
    ]
    for row, col, label in mech_results:
        cell = t9.cell(row, col)
        old_val = cell.text.strip()
        ph = add_field(f'table[9].cell[{row},{col}] (力学-{label})', old_val)
        _replace_all_text_in_container(cell, ph)

    # ============================================================
    # SAVE
    # ============================================================
    doc.save(output_path)

    # Build mapping JSON
    mapping = {
        'total_fields': len(fields),
        'categories': {
            'cover_info': [f for f in fields if 'paragraph[' in f['location'] and
                           any(k in f['location'] for k in
                               ['项目名称', '产品供应商', '生产地址', '电话', '报告日期'])],
            'certificates': [f for f in fields if 'table[' in f['location'] and
                             any(k in f['location'] for k in
                                 ['规格型号', '生产日期', '出厂编号', '检验员',
                                  '检验依据', '检验部门', '检验结果']) and
                             'table[8]' not in f['location'] and
                             'table[9]' not in f['location']],
            'visual': [f for f in fields if '外观' in f['location'] or
                       'table[8]' in f['location']],
            'mechanical': [f for f in fields if '力学' in f['location'] or
                           'table[9]' in f['location']],
        },
    }

    # Catch unclassified fields
    classified = set()
    for cat_items in mapping['categories'].values():
        for f in cat_items:
            classified.add(f['id'])
    mapping['categories']['other'] = [f for f in fields
                                       if f['id'] not in classified]

    # Clean up categories field format for JSON
    for cat in mapping['categories']:
        mapping['categories'][cat] = [
            {'placeholder': f['placeholder'],
             'original': f['original'],
             'location': f['location']}
            for f in mapping['categories'][cat]
        ]

    with open(mapping_path, 'w', encoding='utf-8') as f:
        json.dump(mapping, f, ensure_ascii=False, indent=2)

    # Print summary
    print(f'Total fields: {len(fields)}')
    for cat, items in mapping['categories'].items():
        print(f'  {cat}: {len(items)} fields')

    return mapping


if __name__ == '__main__':
    base = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(base, 'template.docx')
    output_path = os.path.join(base, 'template_prepared.docx')
    mapping_path = os.path.join(base, 'param_mapping.json')

    prepare_template(template_path, output_path, mapping_path)
    print(f'\nSaved: {output_path}')
    print(f'Saved: {mapping_path}')
